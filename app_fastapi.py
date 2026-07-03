# -*- coding: utf-8 -*-
"""
app_fastapi.py — FastAPI backend for the batch document extraction system.

Routes:
  POST   /api/templates/parse                          Parse uploaded Excel template
  POST   /api/templates                                Save confirmed template snapshot
  GET    /api/templates/{template_id}                  Retrieve a template snapshot
  GET    /api/templates                                List all templates

  POST   /api/tasks                                    Create a new task
  POST   /api/tasks/{task_id}/files                    Upload documents to a task
  POST   /api/tasks/{task_id}/start                    Start extraction
  GET    /api/tasks/{task_id}                          Get task status + counts
  GET    /api/tasks/{task_id}/files                    Get per-file status list
  POST   /api/tasks/{task_id}/cancel                   Cancel pending files

  GET    /api/tasks/{task_id}/results                  Paginated result list
  GET    /api/tasks/{task_id}/results/{file_id}        Single file detail
  PATCH  /api/tasks/{task_id}/results/{file_id}        Human-edit field values
  POST   /api/tasks/{task_id}/files/{file_id}/retry    Retry a failed file

  POST   /api/tasks/{task_id}/exports/excel            Export results as XLSX
  POST   /api/tasks/{task_id}/exports/docx             Export results as DOCX

  GET    /health                                       Health check
"""

from __future__ import annotations

import io
from tempfile import template
from typing import Any, Dict, List, Optional

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from langgraph.func import task

from export_service import build_docx, build_excel
from schema import (
    CreateTaskRequest,
    ErrorDetail,
    ExportRequest,
    ExportScope,
    PatchResultRequest,
    ParseTemplateRequest,
    SaveTemplateRequest,
    TemplateField,
)
from task_engine import (
    add_files,
    cancel_task,
    create_task,
    get_file_result,
    get_results,
    get_task,
    list_tasks,
    patch_result,
    retry_file,
    start_task,
)
from template_service import (
    get_template,
    list_templates,
    parse_excel,
    save_template,
    switch_sheet,
)


# ══════════════════════════════════════════════════════════════════════════════
# APP SETUP
# ══════════════════════════════════════════════════════════════════════════════

app = FastAPI(
    title="AI Batch Document Extraction API",
    description="Template-driven batch extraction of structured fields from PDF and DOCX documents.",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _http_error(status: int, code: str, message: str, details: Optional[Dict] = None) -> HTTPException:
    return HTTPException(
        status_code=status,
        detail=ErrorDetail(code=code, message=message, details=details).model_dump(),
    )


def _wrap(func, *args, **kwargs):
    """Call a service function and convert known exceptions to HTTPExceptions."""
    try:
        return func(*args, **kwargs)
    except KeyError as e:
        raise _http_error(404, "NOT_FOUND", str(e))
    except ValueError as e:
        raise _http_error(400, "INVALID_REQUEST", str(e))
    except Exception as e:
        raise _http_error(500, "INTERNAL_ERROR", str(e))


# ══════════════════════════════════════════════════════════════════════════════
# HEALTH
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/health", tags=["System"])
def health():
    from image_enhancer import PRESETS
    return {"status": "ok", "ocr_presets": list(PRESETS.keys())}


# ══════════════════════════════════════════════════════════════════════════════
# TEMPLATE ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/templates/parse", tags=["Templates"])
async def parse_template(
    file: UploadFile = File(...),
    sheet_name: Optional[str] = Form(default=None),
    header_row: int = Form(default=1),
):
    """
    Upload an Excel file and parse it into field definitions.
    Returns upload_id, sheet list, and parsed fields for user review.
    """
    if not file.filename.lower().endswith(".xlsx"):
        raise _http_error(400, "INVALID_FILE_TYPE", "只接受 .xlsx 文件。")
    raw = await file.read()
    result = _wrap(parse_excel, raw, file.filename, sheet_name, header_row)
    return result


@app.post("/api/templates/parse/switch-sheet", tags=["Templates"])
def switch_template_sheet(body: dict):
    """Re-parse the same uploaded file using a different sheet."""
    upload_id  = body.get("upload_id")
    sheet_name = body.get("sheet_name")
    header_row = int(body.get("header_row", 1))
    if not upload_id or not sheet_name:
        raise _http_error(400, "MISSING_PARAMS", "upload_id and sheet_name are required.")
    return _wrap(switch_sheet, upload_id, sheet_name, header_row)


@app.post("/api/templates", tags=["Templates"])
def save_template_route(body: SaveTemplateRequest):
    """Save a user-confirmed template snapshot. Returns template_id."""
    if not body.fields:
        raise _http_error(400, "EMPTY_FIELDS", "模板至少包含一个有效字段。")
    snapshot = _wrap(save_template, body.upload_id, body.fields, body.name)
    return {
        "template_id": snapshot.template_id,
        "name": snapshot.name,
        "version": snapshot.version,
        "field_count": len(snapshot.fields),
        "created_at": snapshot.created_at.isoformat(),
    }


@app.get("/api/templates", tags=["Templates"])
def list_templates_route():
    """List all saved templates, newest first."""
    templates = list_templates()
    return [
        {
            "template_id": t.template_id,
            "name": t.name,
            "version": t.version,
            "field_count": len(t.fields),
            "source_file": t.source_file,
            "created_at": t.created_at.isoformat(),
        }
        for t in templates
    ]


@app.get("/api/templates/{template_id}", tags=["Templates"])
def get_template_route(template_id: str):
    """Retrieve a full template snapshot including all field definitions."""
    return _wrap(get_template, template_id)


# ══════════════════════════════════════════════════════════════════════════════
# TASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/tasks", tags=["Tasks"])
def create_task_route(body: CreateTaskRequest):
    """Create a new draft task linked to a confirmed template."""
    task = _wrap(create_task, body.template_id, body.name, body.settings)
    return _task_summary(task)


@app.get("/api/tasks", tags=["Tasks"])
def list_tasks_route():
    """List all tasks, newest first."""
    return [_task_summary(t) for t in list_tasks()]


@app.get("/api/tasks/{task_id}", tags=["Tasks"])
def get_task_route(task_id: str):
    """Get task status, progress percentage, and per-status counts."""
    task = _wrap(get_task, task_id)
    return _task_summary(task)


@app.post("/api/tasks/{task_id}/files", tags=["Tasks"])
async def upload_files_route(task_id: str, files: List[UploadFile] = File(...)):
    """
    Upload one or more documents to a draft task.
    Each file is validated individually — one bad file does not fail the whole request.
    """
    file_tuples = []
    for uf in files:
        raw = await uf.read()
        file_tuples.append((uf.filename, raw))
    return _wrap(add_files, task_id, file_tuples)


@app.post("/api/tasks/{task_id}/start", tags=["Tasks"])
def start_task_route(task_id: str, background_tasks: BackgroundTasks):
    """
    Start extraction for a draft task.
    Processing runs synchronously in MVP (blocks until complete).
    For production, move _run_extraction to a background worker.
    """
    task = _wrap(start_task, task_id)
    return _task_summary(task)


@app.post("/api/tasks/{task_id}/cancel", tags=["Tasks"])
def cancel_task_route(task_id: str):
    """Cancel all pending files in a task."""
    task = _wrap(cancel_task, task_id)
    return _task_summary(task)


@app.get("/api/tasks/{task_id}/files", tags=["Tasks"])
def get_task_files_route(task_id: str):
    """Get per-file status list for a task."""
    task = _wrap(get_task, task_id)
    return [_file_summary(tf) for tf in task.files]


@app.post("/api/tasks/{task_id}/files/{file_id}/retry", tags=["Tasks"])
def retry_file_route(task_id: str, file_id: str):
    """Retry a single failed or cancelled file."""
    tf = _wrap(retry_file, task_id, file_id)
    return _file_summary(tf)


# ══════════════════════════════════════════════════════════════════════════════
# RESULT ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/tasks/{task_id}/results", tags=["Results"])
def get_results_route(
    task_id: str,
    status:     Optional[str] = Query(default=None),
    keyword:    Optional[str] = Query(default=None),
    has_issues: bool          = Query(default=False),
    page:       int           = Query(default=1, ge=1),
    page_size:  int           = Query(default=50, ge=1, le=200),
):
    """
    Paginated result list with optional filters.
    status: success | needs_review | failed | cancelled
    """
    data = _wrap(get_results, task_id, status, keyword, has_issues, page, page_size)
    return {
        "total":     data["total"],
        "page":      data["page"],
        "page_size": data["page_size"],
        "results":   [_result_summary(r) for r in data["results"]],
    }


@app.get("/api/tasks/{task_id}/results/{file_id}", tags=["Results"])
def get_file_result_route(task_id: str, file_id: str):
    """Get the full extraction detail for a single file."""
    result = _wrap(get_file_result, task_id, file_id)
    return _result_detail(result)


@app.patch("/api/tasks/{task_id}/results/{file_id}", tags=["Results"])
def patch_result_route(task_id: str, file_id: str, body: PatchResultRequest):
    """
    Apply human edits to one file's extracted fields.
    Body: { "fields": { "field_key": "new_value", ... } }
    Validates types and recalculates file status.
    """
    result = _wrap(patch_result, task_id, file_id, body.fields)
    return _result_detail(result)


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/tasks/{task_id}/exports/excel", tags=["Exports"])
def export_excel_route(task_id: str, body: ExportRequest):
    """
    Generate and stream an Excel workbook with extraction results.
    Sheet 1 (提取结果): one row per document, one column per template field.
    Sheet 2 (处理记录): file name, status, issues, timing, model.
    """
    task    = _wrap(get_task, task_id)
    template = _wrap(get_template, task.template_id)
    results = _select_results(task, body)
    from export_service import build_excel, build_docx
    # in export_excel_route:
    xl_bytes = build_excel(task, template, results, body.include_log)

    # in export_docx_route:
    docx_bytes = build_docx(task, template, results, body.include_log)
    filename = f"extraction_{task.task_id[:8]}.xlsx"
    return StreamingResponse(
        io.BytesIO(xl_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/api/tasks/{task_id}/exports/docx", tags=["Exports"])
def export_docx_route(task_id: str, body: ExportRequest):
    """
    Generate and stream a DOCX document.
    One H1 section per document; fields listed in template order.
    Failed files appear in an appendix section.
    """
    task     = _wrap(get_task, task_id)
    template = _wrap(get_template, task.template_id)
    results  = _select_results(task, body)
    docx_bytes = _build_docx_export(task, template, results, body.include_log)
    filename = f"extraction_{task.task_id[:8]}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def _select_results(task, body: ExportRequest):
    from schema import FileStatus
    if body.scope == ExportScope.SELECTED and body.file_ids:
        return [r for r in task.results if r.file_id in body.file_ids]
    if body.scope == ExportScope.ALL:
        return list(task.results)
    # default: success + needs_review
    return [r for r in task.results if r.status in (FileStatus.SUCCESS, FileStatus.NEEDS_REVIEW)]


def _build_excel_export(task, template, results, include_log: bool) -> bytes:
    from datetime import datetime as dt
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.worksheet.table import Table, TableStyleInfo

    HDR_FONT  = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    HDR_FILL  = PatternFill("solid", fgColor="4B2E83")
    BODY_FONT = Font(name="Calibri", size=11)
    CENTER    = Alignment(horizontal="center", vertical="center")
    LEFT      = Alignment(horizontal="left",   vertical="center", indent=1)
    THIN      = Side(style="thin", color="DDDDDD")
    BRD       = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

    wb = Workbook()
    ws = wb.active
    ws.title = "提取结果"

    # Build ordered field list from template
    ordered_fields = sorted(template.fields, key=lambda f: f.order)

    # Header row
    meta_headers = ["文件名", "状态", "问题数"]
    field_headers = [f.name for f in ordered_fields]
    all_headers = meta_headers + field_headers + ["最后修改"]

    for ci, h in enumerate(all_headers, start=1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = HDR_FONT; c.fill = HDR_FILL; c.alignment = CENTER; c.border = BRD
    ws.row_dimensions[1].height = 22

    STATUS_LABELS = {
        "success": "成功", "needs_review": "待检查",
        "failed": "失败", "cancelled": "已取消", "pending": "等待", "processing": "处理中",
    }

    for ri, result in enumerate(results, start=2):
        ws.cell(row=ri, column=1, value=result.filename).font = BODY_FONT
        ws.cell(row=ri, column=2, value=STATUS_LABELS.get(result.status.value, result.status.value)).font = BODY_FONT
        ws.cell(row=ri, column=3, value=result.issue_count).font = BODY_FONT

        for ci, fdef in enumerate(ordered_fields, start=4):
            fv = result.get_field(fdef.key)
            val = fv.value if fv else None
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            c = ws.cell(row=ri, column=ci, value=val)
            c.font = BODY_FONT; c.alignment = LEFT; c.border = BRD
            if fv and fv.manually_edited:
                c.font = Font(name="Calibri", size=11, italic=True, color="1A5276")

        last_col = len(all_headers)
        last_mod = result.completed_at.strftime("%Y-%m-%d %H:%M") if result.completed_at else ""
        ws.cell(row=ri, column=last_col, value=last_mod).font = BODY_FONT
        ws.row_dimensions[ri].height = 18

    # Auto column widths
    for ci in range(1, len(all_headers) + 1):
        max_len = max(
            (len(str(ws.cell(row=r, column=ci).value or "")) for r in range(1, len(results) + 2)),
            default=10,
        )
        from openpyxl.utils import get_column_letter
        ws.column_dimensions[get_column_letter(ci)].width = min(max(max_len + 3, 12), 50)

    # Excel Table
    if results:
        from openpyxl.utils import get_column_letter as gcl
        last_col_letter = gcl(len(all_headers))
        tbl = Table(displayName="ExtractionResults", ref=f"A1:{last_col_letter}{len(results)+1}")
        tbl.tableStyleInfo = TableStyleInfo(name="TableStyleMedium4", showRowStripes=True)
        ws.add_table(tbl)

    ws.freeze_panes = "A2"

    # Sheet 2: processing log
    if include_log:
        ws2 = wb.create_sheet("处理记录")
        log_headers = ["文件名", "状态", "问题数", "错误信息", "开始时间", "完成时间", "耗时(秒)", "模型", "重试次数"]
        for ci, h in enumerate(log_headers, start=1):
            c = ws2.cell(row=1, column=ci, value=h)
            c.font = HDR_FONT; c.fill = HDR_FILL; c.alignment = CENTER
        for ri, result in enumerate(task.results, start=2):
            row_data = [
                result.filename,
                STATUS_LABELS.get(result.status.value, result.status.value),
                result.issue_count,
                result.error_message or "",
                result.started_at.strftime("%Y-%m-%d %H:%M:%S") if result.started_at else "",
                result.completed_at.strftime("%Y-%m-%d %H:%M:%S") if result.completed_at else "",
                round(result.elapsed_seconds, 1) if result.elapsed_seconds else "",
                result.model_used or "",
                result.retry_count,
            ]
            for ci, val in enumerate(row_data, start=1):
                ws2.cell(row=ri, column=ci, value=val).font = BODY_FONT
        for ci in range(1, len(log_headers) + 1):
            from openpyxl.utils import get_column_letter
            ws2.column_dimensions[get_column_letter(ci)].width = 20

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_docx_export(task, template, results, include_log: bool) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from schema import FileStatus

    doc = Document()

    # Title page
    title = doc.add_heading(task.name, level=0)
    from datetime import datetime as dt
    doc.add_paragraph(f"导出时间：{dt.now().strftime('%Y-%m-%d %H:%M')}")
    counts = task.counts
    doc.add_paragraph(
        f"文档总数：{counts['total']}  成功：{counts['success']}  "
        f"待检查：{counts['needs_review']}  失败：{counts['failed']}"
    )
    doc.add_page_break()

    ordered_fields = sorted(template.fields, key=lambda f: f.order)
    ok_results   = [r for r in results if r.status in (FileStatus.SUCCESS, FileStatus.NEEDS_REVIEW)]
    fail_results = [r for r in results if r.status == FileStatus.FAILED]

    STATUS_LABELS = {
        "success": "成功", "needs_review": "待检查",
        "failed": "失败", "cancelled": "已取消",
    }

    for result in ok_results:
        doc.add_heading(result.filename, level=1)
        status_label = STATUS_LABELS.get(result.status.value, result.status.value)
        doc.add_paragraph(f"状态：{status_label}  |  问题数：{result.issue_count}")
        for fdef in ordered_fields:
            fv = result.get_field(fdef.key)
            val = fv.value if fv else None
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            display = str(val) if val is not None else "（未找到）"
            p = doc.add_paragraph()
            run_label = p.add_run(fdef.name + "：")
            run_label.bold = True
            run_val = p.add_run(display)
            if fv and fv.manually_edited:
                run_val.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        doc.add_paragraph()  # spacing between docs

    if include_log and fail_results:
        doc.add_page_break()
        doc.add_heading("处理失败记录", level=1)
        for result in fail_results:
            doc.add_heading(result.filename, level=2)
            doc.add_paragraph(f"错误原因：{result.error_message or '未知错误'}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# RESPONSE SERIALISERS
# ══════════════════════════════════════════════════════════════════════════════

def _task_summary(task) -> Dict[str, Any]:
    return {
        "task_id":      task.task_id,
        "name":         task.name,
        "template_id":  task.template_id,
        "status":       task.status.value,
        "progress_pct": task.progress_pct,
        "counts":       {k: v for k, v in task.counts.items()},
        "settings":     task.settings.model_dump(),
        "created_at":   task.created_at.isoformat(),
        "started_at":   task.started_at.isoformat() if task.started_at else None,
        "completed_at": task.completed_at.isoformat() if task.completed_at else None,
    }


def _file_summary(tf) -> Dict[str, Any]:
    return {
        "file_id":   tf.file_id,
        "filename":  tf.filename,
        "file_type": tf.file_type,
        "size":      tf.size,
        "status":    tf.status.value,
    }


def _result_summary(result) -> Dict[str, Any]:
    return {
        "file_id":      result.file_id,
        "filename":     result.filename,
        "status":       result.status.value,
        "issue_count":  result.issue_count,
        "elapsed_seconds": result.elapsed_seconds,
        "model_used":   result.model_used,
        "retry_count":  result.retry_count,
        "fields": {
            fv.key: {
                "value":           fv.value,
                "manually_edited": fv.manually_edited,
            }
            for fv in result.fields
        },
        "issues": [
            {
                "field_key":  i.field_key,
                "field_name": i.field_name,
                "issue_type": i.issue_type.value,
                "message":    i.message,
            }
            for i in result.issues
        ],
    }


def _result_detail(result) -> Dict[str, Any]:
    base = _result_summary(result)
    base["fields"] = [
        {
            "key":             fv.key,
            "name":            fv.name,
            "value":           fv.value,
            "raw_ai_value":    fv.raw_ai_value,
            "manually_edited": fv.manually_edited,
            "edited_at":       fv.edited_at.isoformat() if fv.edited_at else None,
            "source_snippet":  fv.source_snippet,
        }
        for fv in result.fields
    ]
    return base


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT  (uvicorn app_fastapi:app --reload)
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app_fastapi:app", host="0.0.0.0", port=8000, reload=True)