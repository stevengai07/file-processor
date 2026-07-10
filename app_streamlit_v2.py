# -*- coding: utf-8 -*-
"""
app_streamlit_v2.py — Three-page Streamlit UI for AI Batch Document Extraction.

Pages:
  1. 模板配置   — Upload Excel template, preview/edit fields, confirm snapshot
  2. 批量提取   — Create task, upload documents, configure settings, run extraction
  3. 结果审核   — Review results table, edit fields inline, export Excel/DOCX
  4. AI 控制台  — Single or multi-document interactive extraction, summarization, and export

Run:
  streamlit run app_streamlit_v2.py
"""

from __future__ import annotations

import io
import json
import time
import datetime
import re
import os
import pandas as pd
from typing import Any, Dict, List, Optional

import streamlit as st

# ── page config (must be first Streamlit call) ────────────────────────────
st.set_page_config(
    page_title="AI 批量文档提取",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ══════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS & THEME INJECTION
# ══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

/* ══ DESIGN TOKENS ══ */
:root {
    --brand:        #4B2E83;
    --brand-light:  #7B5EA7;
    --brand-xlight: #EDE7F6;
    --brand-dark:   #2D1B69;
    --surface:      #FFFFFF;
    --surface-soft: #F7F6FB;
    --surface-dim:  #F0EDFA;
    --border:       rgba(75,46,131,0.12);
    --text-primary: #1A1A2E;
    --text-muted:   #4B5563;
    --text-faint:   #9CA3AF;
    --success:      #10B981;
    --warning:      #F59E0B;
    --error:        #EF4444;
    --info:         #3B82F6;
    --radius-sm:    6px;
    --radius-md:    10px;
    --radius-lg:    14px;
    --shadow-sm:    0 1px 2px rgba(45,27,105,0.06), 0 1px 3px rgba(45,27,105,0.1);
    --shadow-md:    0 4px 6px rgba(45,27,105,0.08), 0 2px 4px rgba(45,27,105,0.06);
    --shadow-hover: 0 10px 15px rgba(45,27,105,0.1), 0 4px 6px rgba(45,27,105,0.05);
    --transition:   all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
}

/* ══ BASE APP STYLING ══ */
[data-testid="stAppViewContainer"] {
    background-color: var(--surface-soft);
    font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
}

/* ══ SIDEBAR MAGIC ══ */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1E1145 0%, #2D1B69 50%, #3D2280 100%);
    border-right: 1px solid rgba(255,255,255,0.05);
}
[data-testid="stSidebar"] * {
    color: #F3F4F6 !important;
}
[data-testid="stSidebar"] .stSelectbox > div > div, 
[data-testid="stSidebar"] .stTextInput > div > div > input {
    background: rgba(255,255,255,0.1) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    border-radius: var(--radius-md) !important;
    color: #FFFFFF !important;
}
.prog-bar-wrap {
    background: rgba(255,255,255,0.15);
    border-radius: 999px; height: 8px; overflow: hidden; margin: 10px 0;
    box-shadow: inset 0 1px 2px rgba(0,0,0,0.2);
}
.prog-bar-fill {
    height: 100%; border-radius: 999px;
    background: linear-gradient(90deg, #A78BFA, #60A5FA, #34D399);
    background-size: 200% 200%;
    animation: gradientShift 3s ease infinite;
    transition: width 0.4s ease;
}
@keyframes gradientShift { 0%{background-position:0% 50%} 50%{background-position:100% 50%} 100%{background-position:0% 50%} }

/* ══ PAGE SECTION HEADERS ══ */
.sec-hdr {
    display: flex; align-items: center; gap: 12px;
    font-size: 1.35rem; font-weight: 800; color: var(--brand);
    margin: 0 0 24px 0; padding: 16px 24px;
    background: var(--surface);
    border-radius: var(--radius-lg);
    box-shadow: var(--shadow-sm);
    border-left: 5px solid var(--brand);
    letter-spacing: -0.01em;
}

/* ══ MODERN HTML CARDS ══ */
.card {
    background: var(--surface);
    border-radius: var(--radius-lg);
    padding: 24px;
    box-shadow: var(--shadow-sm);
    border: 1px solid var(--border);
    margin-bottom: 20px;
    transition: var(--transition);
}
.card:hover { box-shadow: var(--shadow-hover); transform: translateY(-2px); }
.card-title {
    font-size: 0.95rem; font-weight: 700; color: var(--brand);
    text-transform: uppercase; letter-spacing: 0.05em;
    margin-bottom: 16px; padding-bottom: 12px;
    border-bottom: 2px solid var(--surface-dim);
    display: flex; align-items: center; gap: 8px;
}

/* ══ NATIVE COMPONENT TWEAKS ══ */
[data-testid="stButton"] button {
    border-radius: var(--radius-md) !important;
    font-weight: 600 !important;
    transition: var(--transition) !important;
}
[data-testid="stButton"] button[kind="primary"] {
    background: linear-gradient(135deg, var(--brand) 0%, var(--brand-light) 100%) !important;
    border: none !important;
    box-shadow: 0 4px 6px rgba(75,46,131,0.2) !important;
}
[data-testid="stButton"] button[kind="primary"]:hover {
    box-shadow: 0 6px 12px rgba(75,46,131,0.3) !important;
    transform: translateY(-1px);
}

/* ══ DATA EDITOR (Page 1 & 3) ══ */
[data-testid="stDataFrame"] {
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border) !important;
    box-shadow: var(--shadow-sm) !important;
}

/* ══ KPI METRIC TILES ══ */
.kpi-row { display: grid; grid-template-columns: repeat(auto-fit, minmax(140px, 1fr)); gap: 16px; margin-bottom: 24px; }
.kpi {
    background: var(--surface); border-radius: var(--radius-lg);
    padding: 20px; text-align: center;
    box-shadow: var(--shadow-sm); border: 1px solid var(--border);
    transition: var(--transition);
}
.kpi:hover { box-shadow: var(--shadow-md); transform: translateY(-3px); border-color: var(--brand-light); }
.kpi-val  { font-size: 2.2rem; font-weight: 800; line-height: 1; letter-spacing: -0.03em; margin-bottom: 8px; }
.kpi-lbl  { font-size: 0.75rem; color: var(--text-muted); font-weight: 600; text-transform: uppercase; letter-spacing: 0.05em; }
.kpi-blue  { color: #2563EB; }
.kpi-green { color: #059669; }
.kpi-amber { color: #D97706; }
.kpi-red   { color: #DC2626; }
.kpi-grey  { color: #4B5563; }

/* ══ BADGES & CHIPS ══ */
.badge {
    display: inline-flex; align-items: center; gap: 4px;
    border-radius: 999px; padding: 4px 12px;
    font-size: 0.75rem; font-weight: 700; white-space: nowrap;
}
.badge-success  { background: #D1FAE5; color: #065F46; }
.badge-review   { background: #FEF3C7; color: #92400E; }
.badge-failed   { background: #FEE2E2; color: #991B1B; }

/* ══ FILE ROW HOVER ══ */
.file-row {
    display: flex; align-items: center; gap: 12px;
    padding: 12px 16px; background: var(--surface);
    border-radius: var(--radius-md); margin-bottom: 8px;
    border: 1px solid var(--border);
    transition: var(--transition);
    cursor: default;
}
.file-row:hover { border-color: var(--brand-light); background: var(--surface-soft); box-shadow: var(--shadow-sm); }
.file-name { flex: 1; font-size: 0.9rem; font-weight: 500; color: var(--text-primary); }

/* ══ FIELD GRID FOR RESULTS ══ */
.field-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(280px, 1fr)); gap: 12px; }
.field-item {
    background: #F9FAFB; border: 1px solid #E5E7EB;
    border-radius: var(--radius-md); padding: 14px 18px;
    transition: var(--transition);
}
.field-item:hover { border-color: var(--brand-light); background: #FFFFFF; box-shadow: var(--shadow-sm); }
.field-item-label {
    font-size: 0.75rem; color: var(--brand); font-weight: 700;
    text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 6px;
}
.field-item-value { font-size: 0.95rem; color: #111827; word-break: break-word; line-height: 1.5; }
.field-item-edited { border-left: 4px solid var(--info); background: #EFF6FF; border-color: #BFDBFE; }
.edited-dot { color: var(--info); font-size: 0.7rem; margin-left: 6px; font-weight: 600; padding: 2px 6px; background: #DBEAFE; border-radius: 4px;}

</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INITIALISATION
# ══════════════════════════════════════════════════════════════════════════════

def _init_state():
    defaults = {
        "page":           "1_template",
        "upload_id":      None,
        "parse_resp":     None,
        "template_id":    None,
        "template_name":  "",
        "task_id":        None,
        "task":           None,
        "results":        [],
        "selected_file":  None,
        "edit_mode":      False,
        "export_log":     [],
        # AI 控制台
        "chat_history":   [],
        "chat_doc_text":  "",
        "chat_doc_name":  "",
        "chat_template_text": "",
        "console_raw":    None,
        "console_raw_name": "",
        "console_file_list": [],
        "console_tpl_list": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════

def _sidebar():
    with st.sidebar:
        st.markdown("""
<div style="padding:16px 4px 12px;margin-bottom:4px;">
  <div style="display:flex;align-items:center;gap:12px;margin-bottom:6px;">
    <div style="width:40px;height:40px;background:linear-gradient(135deg,#7B5EA7,#A78BFA);border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:20px;box-shadow:0 4px 10px rgba(0,0,0,0.3);">📄</div>
    <div>
      <div style="font-size:1.1rem;font-weight:800;color:#FFFFFF;letter-spacing:-0.01em;line-height:1.2;">AI 批量提取</div>
      <div style="font-size:0.65rem;color:rgba(233,227,247,0.6);font-weight:600;letter-spacing:0.06em;">DOCUMENT EXTRACTION</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)
        st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.1);margin:0 0 16px;'>", unsafe_allow_html=True)

        pages = {
            "1_template": "① 模板配置",
            "2_extract":  "② 批量提取",
            "3_review":   "③ 结果审核",
            "4_console":  "④ AI 控制台",
        }
        for key, label in pages.items():
            active = st.session_state.page == key
            if st.button(
                label,
                key=f"nav_{key}",
                use_container_width=True,
                type="primary" if active else "secondary",
            ):
                st.session_state.page = key
                st.rerun()

        st.markdown("---")

        page_help = {
            "1_template": {
                "icon": "📋", "title": "① 模板配置",
                "用途": "定义 AI 需提取哪些字段（如金额、日期）。",
                "格式要求": "上传 Excel（.xlsx），首行为字段名。",
            },
            "2_extract": {
                "icon": "⚡", "title": "② 批量提取",
                "用途": "批量上传 PDF/Word，AI 自动生成结构化表格。",
                "格式要求": "PDF 或 DOCX，单文件建议 ≤ 20MB。",
            },
            "3_review": {
                "icon": "✅", "title": "③ 结果审核",
                "用途": "在线人工复核与修改，一键导出 Excel / Word。",
                "格式要求": "数据承接自上一步，无需重新上传。",
            },
            "4_console": {
                "icon": "🧠", "title": "④ AI 控制台",
                "用途": "多文档跨篇章追问，基于样板仿写汇总与公文生成。",
                "格式要求": "支持上传【样板】及【多份目标文件】。",
            },
        }
        current_page = st.session_state.get("page", "1_template")
        if current_page in page_help:
            h = page_help[current_page]
            st.markdown(
                f"""<div style="
                    background: rgba(255,255,255,0.06);
                    border-radius: 12px;
                    padding: 16px;
                    margin-bottom: 8px;
                    border: 1px solid rgba(255,255,255,0.1);
                ">
                <div style="font-size:1rem;font-weight:700;color:#EDE7F6;margin-bottom:8px;">
                    {h['icon']} {h['title']}
                </div>
                <div style="font-size:0.8rem;color:#C5B8E8;line-height:1.6;">
                    <span style="color:#E8DEF8;font-weight:600;">🎯 用途：</span>{h['用途']}<br>
                    <span style="color:#E8DEF8;font-weight:600;">📁 要求：</span>{h['格式要求']}
                </div>
                </div>""",
                unsafe_allow_html=True,
            )

        st.markdown("---")

        tmpl_ok = st.session_state.template_id is not None
        task_ok = st.session_state.task_id is not None
        tmpl_color = "#34D399" if tmpl_ok else "rgba(255,255,255,0.2)"
        task_color = "#34D399" if task_ok else "rgba(255,255,255,0.2)"
        tmpl_label = st.session_state.template_name[:16] + ("…" if len(st.session_state.template_name) > 16 else "") if tmpl_ok else "未配置"
        task_label = (st.session_state.task_id[:8] + "…") if task_ok else "未创建"
        
        st.markdown(f"""
<div style="display:flex;flex-direction:column;gap:8px;margin-bottom:12px;">
  <div style="display:flex;align-items:center;gap:10px;background:rgba(0,0,0,0.15);border-radius:8px;padding:10px 12px; border:1px solid rgba(255,255,255,0.05);">
    <div style="width:10px;height:10px;border-radius:50%;background:{tmpl_color};flex-shrink:0;box-shadow:0 0 8px {tmpl_color};"></div>
    <div style="flex:1;min-width:0;">
      <div style="font-size:0.65rem;color:rgba(233,227,247,0.6);font-weight:700;text-transform:uppercase;letter-spacing:0.05em;">当前模板</div>
      <div style="font-size:0.85rem;color:#E9E3F7;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{tmpl_label}</div>
    </div>
  </div>
  <div style="display:flex;align-items:center;gap:10px;background:rgba(0,0,0,0.15);border-radius:8px;padding:10px 12px; border:1px solid rgba(255,255,255,0.05);">
    <div style="width:10px;height:10px;border-radius:50%;background:{task_color};flex-shrink:0;box-shadow:0 0 8px {task_color};"></div>
    <div style="flex:1;min-width:0;">
      <div style="font-size:0.65rem;color:rgba(233,227,247,0.6);font-weight:700;text-transform:uppercase;letter-spacing:0.05em;">当前任务</div>
      <div style="font-size:0.85rem;color:#E9E3F7;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{task_label}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

        results = st.session_state.results
        if results:
            from collections import Counter
            c = Counter(r.status.value for r in results)
            total = len(results)
            done  = c.get("success", 0) + c.get("needs_review", 0) + c.get("failed", 0) + c.get("cancelled", 0)
            pct   = int(done / total * 100) if total else 0
            st.markdown(f"<div style='font-size:0.85rem;font-weight:600;margin-top:8px;'>任务进度: {done}/{total} ({pct}%)</div>", unsafe_allow_html=True)
            st.markdown(
                f"<div class='prog-bar-wrap'><div class='prog-bar-fill' style='width:{pct}%'></div></div>",
                unsafe_allow_html=True,
            )

        st.markdown("---")

        st.markdown("#### 🌐 语言")
        lang_options = {
            "中文 (Chinese)":  "chi_sim+eng",
            "English":         "eng",
            "日本語 (Japanese)": "jpn+eng",
            "한국어 (Korean)":   "kor+eng",
            "繁體中文 (Traditional)": "chi_tra+eng",
        }
        selected_lang_label = st.selectbox(
            "",
            list(lang_options.keys()),
            index=0,
            key="sidebar_lang",
            label_visibility="collapsed",
        )
        st.session_state["ocr_lang_resolved"] = lang_options[selected_lang_label]

        st.markdown("---")

        st.markdown("#### ⊕ 模型")
        provider_map = {
            "OpenAI":    ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "o1-mini", "o3-mini"],
            "Anthropic": ["claude-3-5-haiku-20241022", "claude-3-5-sonnet-20241022", "claude-3-opus-20240229"],
            "Alibaba":   ["qwen-turbo", "qwen-plus", "qwen-max"],
            "DeepSeek":  ["deepseek-chat", "deepseek-coder"],
            "Gemini":    ["gemini-1.5-flash", "gemini-1.5-pro"],
        }
        selected_provider = st.selectbox("", list(provider_map.keys()), index=0, key="sidebar_provider", label_visibility="collapsed")
        model_list = provider_map[selected_provider]
        selected_model = st.selectbox("", model_list, index=0, key="sidebar_model", label_visibility="collapsed")
        st.session_state["selected_model"] = selected_model

        st.markdown("---")

        st.markdown("#### 🖼️ 扫描文件")
        ocr_preset = st.selectbox("", ["扫描仪", "照片", "混合", "关闭"], index=0, key="sidebar_ocr_preset", label_visibility="collapsed")
        preset_map = {"扫描仪": "scanner", "照片": "photo", "混合": "mixed", "关闭": "off"}
        st.session_state["ocr_preset_resolved"] = preset_map[ocr_preset]

        st.markdown("---")

        provider_key_labels = {
            "OpenAI":    ("OPENAI_API_KEY",    "sk-..."),
            "Anthropic": ("ANTHROPIC_API_KEY", "sk-ant-..."),
            "Alibaba":   ("DASHSCOPE_API_KEY", "sk-...（阿里云 Dashscope）"),
            "DeepSeek":  ("DEEPSEEK_API_KEY",  "sk-...（DeepSeek）"),
            "Gemini":    ("GOOGLE_API_KEY",     "AIza...（Google AI Studio）"),
        }
        env_var, placeholder = provider_key_labels.get(selected_provider, ("API_KEY", "sk-..."))
        key_already_set = bool(os.environ.get(env_var, ""))
        expander_label = f"🔑 API 密钥{'  ✅' if key_already_set else '  ⚠️'}"
        
        with st.expander(expander_label):
            api_key_input = st.text_input(
                f"{env_var}",
                type="password",
                placeholder=placeholder,
            )
            if api_key_input:
                os.environ[env_var] = api_key_input
                st.toast(f"✅ {env_var} 设置成功！", icon="🔑")
                st.rerun()

        st.markdown("---")
        st.caption("Powered by LangChain + OpenAI / Anthropic")


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — TEMPLATE CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

def page_template():
    st.markdown("<div class='sec-hdr'>📋 模板配置 <span style='font-size:0.8rem;font-weight:600;color:#9CA3AF;margin-left:auto;background:#F3F4F6;padding:4px 10px;border-radius:20px;'>Step 1 of 4</span></div>", unsafe_allow_html=True)

    with st.container(border=True):
        st.markdown("<div class='card-title'>📂 上传模板文件</div>", unsafe_allow_html=True)
        col1, col2 = st.columns([3, 1])
        with col1:
            uploaded = st.file_uploader(
                "选择 Excel 模板文件（.xlsx）",
                type=["xlsx"],
                label_visibility="collapsed",
            )
        with col2:
            header_row = st.number_input("表头行号", min_value=1, max_value=10, value=1, help="Excel中字段名称所在的行数")

    if uploaded:
        raw = uploaded.read()
        with st.spinner("解析模板中…"):
            from template_service import parse_excel
            resp = parse_excel(raw, uploaded.name, header_row=header_row)
        st.session_state.upload_id  = resp.upload_id
        st.session_state.parse_resp = resp

        if resp.errors:
            for err in resp.errors:
                st.toast(f"⚠️ {err}", icon="⚠️")

    if st.session_state.parse_resp is None:
        _template_hint()
        return

    resp = st.session_state.parse_resp

    if len(resp.sheets) > 1:
        with st.container(border=True):
            st.markdown("<div class='card-title'>📋 选择工作表</div>", unsafe_allow_html=True)
            chosen = st.selectbox("工作表", resp.sheets, index=resp.sheets.index(resp.selected_sheet))
            if chosen != resp.selected_sheet and st.button("切换工作表"):
                from template_service import switch_sheet
                resp = switch_sheet(st.session_state.upload_id, chosen, header_row)
                st.session_state.parse_resp = resp
                st.rerun()

    with st.container(border=True):
        st.markdown("<div class='card-title'>🔍 字段预览与快捷编辑</div>", unsafe_allow_html=True)
        
        if not resp.fields:
            st.error("未解析到有效字段，请检查模板格式。")
            return

        col_req, col_total = st.columns([1, 3])
        col_total.metric("识别字段总数", len(resp.fields))
        col_req.metric("必填字段数", sum(1 for f in resp.fields if f.required))

        st.caption("您可以直接在下方表格中双击单元格修改字段名称、类型、必填项和提示词。支持增加或删除行。")

        type_opts = ["text", "long_text", "integer", "decimal", "date", "boolean", "list"]
        
        df_data = []
        for fdef in resp.fields:
            df_data.append({
                "字段名称": fdef.name,
                "类型": fdef.type.value if fdef.type.value in type_opts else "text",
                "必填": fdef.required,
                "示例值": fdef.example or "",
                "提取提示": fdef.prompt_hint or "",
                "描述": fdef.description or ""
            })
        df = pd.DataFrame(df_data)

        edited_df = st.data_editor(
            df,
            column_config={
                "字段名称": st.column_config.TextColumn("字段名称", required=True),
                "类型": st.column_config.SelectboxColumn("数据类型", options=type_opts, required=True),
                "必填": st.column_config.CheckboxColumn("是否必填", default=False),
                "示例值": st.column_config.TextColumn("示例参考"),
                "提取提示": st.column_config.TextColumn("AI 提取提示"),
                "描述": st.column_config.TextColumn("内部描述")
            },
            use_container_width=True,
            num_rows="dynamic",
            key="template_data_editor"
        )

    with st.container(border=True):
        st.markdown("<div class='card-title'>💾 确认并保存模板</div>", unsafe_allow_html=True)
        tmpl_name = st.text_input("命名此模板配置", value=st.session_state.template_name or uploaded.name.replace(".xlsx", ""))
        col_save, col_clear = st.columns([1, 3])

        if col_save.button("✅ 保存模板并继续", type="primary", use_container_width=True):
            if edited_df.empty:
                st.error("没有可保存的字段。")
            else:
                from schema import FieldType, TemplateField
                import re
                
                final_fields = []
                for i, row in edited_df.iterrows():
                    name = str(row["字段名称"]).strip()
                    if not name: continue
                    key = re.sub(r"[^\w\u4e00-\u9fff]+", "_", name.lower()).strip("_") or f"field_{i}"
                    final_fields.append(TemplateField(
                        key=key, 
                        name=name, 
                        type=FieldType(row["类型"]),
                        required=bool(row["必填"]), 
                        example=str(row["示例值"]) if pd.notna(row["示例值"]) else None,
                        prompt_hint=str(row["提取提示"]) if pd.notna(row["提取提示"]) else None,
                        description=str(row["描述"]) if pd.notna(row["描述"]) else None,
                        order=i,
                    ))

                from template_service import save_template
                snap = save_template(st.session_state.upload_id, final_fields, name=tmpl_name)
                st.session_state.template_id   = snap.template_id
                st.session_state.template_name = snap.name
                st.toast(f"模板「{snap.name}」已保存成功！", icon="🎉")
                time.sleep(0.6)
                st.session_state.page = "2_extract"
                st.rerun()

        if col_clear.button("🗑 清除重新上传", use_container_width=True):
            st.session_state.parse_resp = None
            st.session_state.upload_id  = None
            st.rerun()


def _make_sample_xlsx() -> bytes:
    import io
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "字段模板"
        headers = ["字段名称", "数据类型", "是否必填", "示例值", "提取提示"]
        header_fill = PatternFill(start_color="4B2E83", end_color="4B2E83", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
        sample_rows = [
            ["合同编号",   "text",    "是", "HT-2025-001",      "通常在文件首页或页眉"],
            ["签署日期",   "date",    "是", "2025-06-01",       "格式：YYYY-MM-DD"],
            ["甲方名称",   "text",    "是", "某某科技有限公司",   "合同甲方全称"],
            ["乙方名称",   "text",    "是", "某某服务有限公司",   "合同乙方全称"],
            ["合同金额",   "number",  "否", "100000",           "单位：人民币元"],
            ["合同期限",   "text",    "否", "12个月",           "起止日期或期限描述"],
        ]
        for row_data in sample_rows:
            ws.append(row_data)
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) * 2 for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError:
        import csv
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["字段名称","数据类型","是否必填","示例值","提取提示"])
        writer.writerow(["合同编号","text","是","HT-2025-001","通常在文件首页"])
        return buf.getvalue().encode("utf-8-sig")


def _template_hint():
    with st.container(border=True):
        st.markdown("<div class='card-title'>📥 下载样例模板</div>", unsafe_allow_html=True)
        st.markdown("<p style='font-size:0.9rem;color:#6B7280;margin-bottom:12px;'>下载预填好字段的 Excel 样例，直接修改后上传即可快速开始。字段名、类型、示例值均已填写，支持合同、发票、报告等场景。</p>", unsafe_allow_html=True)
        sample_bytes = _make_sample_xlsx()
        st.download_button(
            label="⬇ 下载标准提取模板 (.xlsx)",
            data=sample_bytes,
            file_name="sample_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    with st.container(border=True):
        st.markdown("<div class='card-title'>💡 模板格式说明</div>", unsafe_allow_html=True)
        st.markdown("""
        <table style='font-size:0.9rem;border-collapse:collapse;width:100%;border:1px solid #E5E7EB; border-radius:8px; overflow:hidden;'>
        <tr style='background:#F3F4F6;'><th style='padding:10px 14px;text-align:left;border-bottom:1px solid #E5E7EB;'>列名</th><th style='padding:10px 14px;text-align:left;border-bottom:1px solid #E5E7EB;'>说明</th><th style='padding:10px 14px;text-align:left;border-bottom:1px solid #E5E7EB;'>同义词识别</th></tr>
        <tr><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'><b>字段名称</b></td><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'>要提取的字段显示名</td><td style='padding:8px 14px;color:#6B7280;border-bottom:1px solid #E5E7EB;'>name / field name / 名称</td></tr>
        <tr style='background:#FAFAFA;'><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'><b>数据类型</b></td><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'>text / integer / date 等</td><td style='padding:8px 14px;color:#6B7280;border-bottom:1px solid #E5E7EB;'>type / 类型 / 字段类型</td></tr>
        <tr><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'><b>是否必填</b></td><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'>是 / yes / true</td><td style='padding:8px 14px;color:#6B7280;border-bottom:1px solid #E5E7EB;'>required / 必填 / mandatory</td></tr>
        <tr style='background:#FAFAFA;'><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'><b>示例值</b></td><td style='padding:8px 14px;border-bottom:1px solid #E5E7EB;'>帮助 AI 理解数据格式</td><td style='padding:8px 14px;color:#6B7280;border-bottom:1px solid #E5E7EB;'>example / sample / 示例</td></tr>
        <tr><td style='padding:8px 14px;'><b>提取提示</b></td><td style='padding:8px 14px;'>补充上下文位置或搜索规则</td><td style='padding:8px 14px;color:#6B7280;'>hint / prompt / 提示</td></tr>
        </table>
        """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — BATCH EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def page_extract():
    st.markdown("<div class='sec-hdr'>⚡ 批量提取 <span style='font-size:0.8rem;font-weight:600;color:#9CA3AF;margin-left:auto;background:#F3F4F6;padding:4px 10px;border-radius:20px;'>Step 2 of 4</span></div>", unsafe_allow_html=True)

    if not st.session_state.template_id:
        st.error("⚠️ **尚未配置模板**: 请先完成第①步模板配置，定义需要提取的字段，然后才能开始批量提取。", icon="⚠️")
        if st.button("📋 前往配置模板", type="primary"):
            st.session_state.page = "1_template"
            st.rerun()
        return

    with st.container(border=True):
        st.markdown("<div class='card-title'>⚙️ 提取引擎配置</div>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)

        task_name   = col1.text_input("任务名称", value=f"任务_{time.strftime('%m%d_%H%M')}")
        _sidebar_model  = st.session_state.get("selected_model", "gpt-4o-mini")
        _sidebar_preset = st.session_state.get("ocr_preset_resolved", "scanner")
        _sidebar_lang   = st.session_state.get("ocr_lang_resolved", "chi_sim+eng")
        concurrency = col2.slider("并发处理数", min_value=1, max_value=10, value=4, help="并发数量越高提取越快，但可能会触发 API 速率限制。")
        col3.markdown(
            f"<div style='font-size:0.85rem;color:#7B5EA7;padding:28px 12px; background:#F0EDFA; border-radius:8px; text-align:center; font-weight:600;'>"
            f"🤖 {_sidebar_model} &nbsp;•&nbsp; 📷 {_sidebar_preset} &nbsp;•&nbsp; 🌐 {_sidebar_lang}"
            f"</div>",
            unsafe_allow_html=True,
        )

        image_enhance = st.checkbox(
            "🖼️ 启用高级图像增强 (倾斜矫正 / 降噪 / 锐化)",
            value=False,
            help="开启后会对扫描件进行物理矫正，显著提升 OCR 识别率，但会增加单张耗时。",
        )
        model      = _sidebar_model
        ocr_preset = _sidebar_preset
        ocr_lang   = _sidebar_lang

    with st.container(border=True):
        st.markdown("<div class='card-title'>📁 上传源文档</div>", unsafe_allow_html=True)

        uploaded_docs = st.file_uploader(
            "选择 PDF 或 DOCX 文件（支持多选拖拽）",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )

        if uploaded_docs:
            st.markdown(f"<div style='margin-top:12px; font-weight:600; color:#4B2E83;'>已装载 {len(uploaded_docs)} 个待处理文件：</div>", unsafe_allow_html=True)
            for uf in uploaded_docs:
                size_kb = len(uf.read()) // 1024
                uf.seek(0)
                st.markdown(
                    f"<div class='file-row'>"
                    f"<span class='file-name'>📄 {uf.name}</span>"
                    f"<span style='color:#9CA3AF;font-size:0.85rem;font-weight:600;'>{size_kb} KB</span>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

        col_start, col_space = st.columns([1, 3])
        start_clicked = col_start.button(
            "🚀 立即开始提取",
            type="primary",
            use_container_width=True,
            disabled=not uploaded_docs,
        )
        if not uploaded_docs:
            col_space.markdown(
                "<div style='padding-top:10px;font-size:0.85rem;color:#9CA3AF;'>← 请先上传至少一个文档</div>",
                unsafe_allow_html=True,
            )

        if start_clicked and uploaded_docs:
            from schema import ExtractionSettings
            from task_engine import add_files, create_task, start_task

            settings = ExtractionSettings(
                model=model,
                ocr_preset=ocr_preset,
                ocr_lang=ocr_lang,
                concurrency=concurrency,
                image_enhance=image_enhance,
            )

            with st.spinner("正在初始化任务队列…"):
                task = create_task(st.session_state.template_id, name=task_name, settings=settings)
                st.session_state.task_id = task.task_id
                st.session_state.task    = task

                file_tuples = []
                for uf in uploaded_docs:
                    uf.seek(0)
                    file_tuples.append((uf.name, uf.read()))

            with st.spinner("文件上云中…"):
                upload_result = add_files(task.task_id, file_tuples)

            accepted = upload_result["accepted"]
            rejected = upload_result["rejected"]

            if rejected:
                for r in rejected:
                    st.toast(f"拒绝文件 {r['filename']}: {r['reason']}", icon="🚫")

            if not accepted:
                st.error("没有可处理的文件，请检查文件格式。")
                return

            progress_placeholder = st.empty()
            with progress_placeholder.container():
                prog_bar = st.progress(0, text="✨ 引擎预热中…")

            start_time = time.time()

            def _run_with_progress():
                import threading
                done_event = threading.Event()

                def _run():
                    start_task(task.task_id)
                    done_event.set()

                t = threading.Thread(target=_run, daemon=True)
                t.start()

                from task_engine import get_task as _get_task
                while not done_event.wait(timeout=0.3):
                    current = _get_task(task.task_id)
                    done_count = sum(
                        1 for f in current.files
                        if f.status.value in ("success", "needs_review", "failed", "cancelled")
                    )
                    total = len(current.files)
                    pct = int(done_count / total * 100) if total else 0
                    prog_bar.progress(pct, text=f"⚡ 正在超并发提取中 ({done_count}/{total})…")

                t.join()
                prog_bar.progress(100, text="🎉 提取全部完成！")

            try:
                _run_with_progress()
            except Exception as e:
                st.error(f"提取引擎中断：{e}")
                return

            elapsed = time.time() - start_time

            from task_engine import get_task as _get_task, get_results
            final_task = _get_task(task.task_id)
            st.session_state.task = final_task

            result_data = get_results(task.task_id, page_size=500)
            st.session_state.results = result_data["results"]

            progress_placeholder.empty()

            from collections import Counter
            c = Counter(r.status.value for r in st.session_state.results)
            total = len(st.session_state.results)

            st.toast("批处理任务已完成！", icon="✅")

            st.markdown("<div class='kpi-row'>" +
                _kpi(total, "文档总数", "blue") +
                _kpi(c.get("success",0), "成功解析", "green") +
                _kpi(c.get("needs_review",0), "建议人工复核", "amber") +
                _kpi(c.get("failed",0), "解析失败", "red") +
                f"<div class='kpi'><div class='kpi-val kpi-grey'>{elapsed:.1f}<span style='font-size:1rem'>s</span></div><div class='kpi-lbl'>引擎总耗时</div></div>" +
                "</div>", unsafe_allow_html=True)

            if st.button("📊 前往结果审核与导出", type="primary", use_container_width=True):
                st.session_state.page = "3_review"
                st.rerun()


def _kpi(val, label, color):
    return f"<div class='kpi'><div class='kpi-val kpi-{color}'>{val}</div><div class='kpi-lbl'>{label}</div></div>"


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — RESULTS REVIEW
# ══════════════════════════════════════════════════════════════════════════════

def page_review():
    st.markdown("<div class='sec-hdr'>✅ 结果审核 <span style='font-size:0.8rem;font-weight:600;color:#9CA3AF;margin-left:auto;background:#F3F4F6;padding:4px 10px;border-radius:20px;'>Step 3 of 4</span></div>", unsafe_allow_html=True)

    results = st.session_state.results
    if not results:
        st.info("💡 当前内存中暂无提取结果。请先完成第②步的文档批量提取。")
        if st.button("🚀 前往批量提取"):
            st.session_state.page = "2_extract"
            st.rerun()
        return

    from collections import Counter
    c = Counter(r.status.value for r in results)
    total = len(results)

    st.markdown("<div class='kpi-row'>" +
        _kpi(total, "文档总数", "blue") +
        _kpi(c.get("success",0), "完美提取", "green") +
        _kpi(c.get("needs_review",0), "置信度低(建议复核)", "amber") +
        _kpi(c.get("failed",0), "提取失败", "red") +
        _kpi(sum(1 for r in results if any(fv.manually_edited for fv in r.fields)), "人工修正数", "grey") +
        "</div>", unsafe_allow_html=True)

    with st.expander("🔎 高级筛选器", expanded=False):
        fc1, fc2, fc3 = st.columns(3)
        filter_status  = fc1.selectbox("按状态过滤", ["全部", "success", "needs_review", "failed", "cancelled"])
        filter_keyword = fc2.text_input("搜寻文件名")
        filter_issues  = fc3.checkbox("仅显示包含警告的文件")

    filtered = results
    if filter_status != "全部":
        filtered = [r for r in filtered if r.status.value == filter_status]
    if filter_keyword:
        kw = filter_keyword.lower()
        filtered = [r for r in filtered if kw in r.filename.lower()]
    if filter_issues:
        filtered = [r for r in filtered if r.issue_count > 0]

    st.caption(f"当前视图显示 **{len(filtered)}** / {total} 份文档。点击左侧列表查看结构化详情。")

    _export_bar()

    list_col, detail_col = st.columns([1, 2.5])

    with list_col:
        with st.container(border=True):
            st.markdown("<div class='card-title'>📋 文件清单</div>", unsafe_allow_html=True)
            STATUS_BADGE = {
                "success":      "<span class='badge badge-success'>✅ 成功</span>",
                "needs_review": "<span class='badge badge-review'>⚠️ 待复核</span>",
                "failed":       "<span class='badge badge-failed'>❌ 失败</span>",
                "cancelled":    "<span class='badge badge-cancel'>⏹ 取消</span>",
                "pending":      "<span class='badge badge-pending'>⏳ 队列中</span>",
            }
            for result in filtered:
                badge = STATUS_BADGE.get(result.status.value, "")
                issue_txt = f" <span style='color:#D97706;font-size:0.75rem;margin-left:4px;'>({result.issue_count} 处警告)</span>" if result.issue_count else ""
                is_sel = st.session_state.selected_file == result.file_id

                if st.button(
                    f"{'👉 ' if is_sel else ''}{result.filename[:25]}...",
                    key=f"sel_{result.file_id}",
                    use_container_width=True,
                    type="primary" if is_sel else "secondary",
                ):
                    st.session_state.selected_file = result.file_id
                    st.session_state.edit_mode = False
                    st.rerun()

                st.markdown(f"<div style='margin-bottom:12px;'>{badge}{issue_txt}</div>", unsafe_allow_html=True)

    with detail_col:
        if st.session_state.selected_file:
            result = next((r for r in results if r.file_id == st.session_state.selected_file), None)
            if result:
                _result_detail_panel(result)
        else:
            st.markdown("""
            <div style="background:#FFFFFF;border:1px dashed rgba(75,46,131,0.2);border-radius:14px;padding:80px 32px;text-align:center;">
              <div style="font-size:3rem;margin-bottom:16px;opacity:0.3;">🔍</div>
              <div style="font-weight:700;color:#6B7280;font-size:1.1rem;margin-bottom:8px;">等待选择</div>
              <div style="color:#9CA3AF;font-size:0.9rem;">请点击左侧文件列表查看 AI 解析详情与源文档片段。</div>
            </div>
            """, unsafe_allow_html=True)


def _result_detail_panel(result):
    from task_engine import patch_result, retry_file

    STATUS_BADGE = {
        "success":      "<span class='badge badge-success'>✅ 解析成功</span>",
        "needs_review": "<span class='badge badge-review'>⚠️ 触发异常预警</span>",
        "failed":       "<span class='badge badge-failed'>❌ 引擎报错</span>",
        "cancelled":    "<span class='badge badge-cancel'>⏹ 任务取消</span>",
    }

    with st.container(border=True):
        st.markdown(
            f"<div style='font-size:1.1rem; font-weight:700; color:#2D1B69; margin-bottom:8px;'>📄 {result.filename}</div>"
            f"<div style='display:flex; align-items:center; gap:12px; margin-bottom:4px;'>"
            + STATUS_BADGE.get(result.status.value, "")
            + (f"<span style='color:#6B7280;font-size:0.85rem;'>⏱ 耗时: <b>{result.elapsed_seconds:.1f}s</b></span>" if result.elapsed_seconds else "")
            + (f"<span style='color:#6B7280;font-size:0.85rem;'>🤖 驱动: <b>{result.model_used}</b></span>" if result.model_used else "")
            + "</div>",
            unsafe_allow_html=True,
        )

    if result.issues:
        with st.container(border=True):
            st.markdown("<div class='card-title' style='color:#D97706;'>⚠️ 异常报告清单</div>", unsafe_allow_html=True)
            ISSUE_LABELS = {
                "required_missing": "必填缺失",
                "type_invalid":     "格式非法",
                "value_ambiguous":  "指代模糊",
                "source_not_found": "未找到溯源",
                "extraction_error": "提取报错",
            }
            for issue in result.issues:
                st.markdown(
                    f"<div class='issue-row'>"
                    f"<span style='font-weight:700;'>[{issue.field_name}]</span> "
                    f"<span style='background:#FDE68A;color:#92400E;padding:2px 6px;border-radius:4px;font-size:0.75rem;margin:0 6px;'>{ISSUE_LABELS.get(issue.issue_type.value, issue.issue_type.value)}</span>"
                    f" {issue.message}"
                    f"</div>",
                    unsafe_allow_html=True,
                )

    if result.fields:
        with st.container(border=True):
            col_view, col_edit_btn = st.columns([4, 1])
            col_view.markdown("<div class='card-title'>📊 结构化字段萃取结果</div>", unsafe_allow_html=True)
            if col_edit_btn.button(
                "✏️ 修正内容" if not st.session_state.edit_mode else "👁 取消编辑",
                key="toggle_edit",
                use_container_width=True
            ):
                st.session_state.edit_mode = not st.session_state.edit_mode
                st.rerun()

            if not st.session_state.edit_mode:
                html_cards = "<div class='field-grid'>"
                for fv in result.fields:
                    val = fv.value
                    if isinstance(val, list):
                        val = ", ".join(str(v) for v in val)
                    elif isinstance(val, bool):
                        val = "是" if val else "否"
                    display = str(val) if val is not None else "<span style='color:#D1D5DB;font-style:italic;'>（未抓取到内容）</span>"
                    edit_cls = " field-item-edited" if fv.manually_edited else ""
                    edit_dot = "<span class='edited-dot'>✏️ 人工修正</span>" if fv.manually_edited else ""
                    html_cards += (
                        f"<div class='field-item{edit_cls}'>"
                        f"<div class='field-item-label'>{fv.name}{edit_dot}</div>"
                        f"<div class='field-item-value'>{display}</div>"
                        f"</div>"
                    )
                html_cards += "</div>"
                st.markdown(html_cards, unsafe_allow_html=True)

            else:
                edited_vals: Dict[str, Any] = {}
                for fv in result.fields:
                    val = fv.value
                    if isinstance(val, list):
                        val = ", ".join(str(v) for v in val)
                    elif isinstance(val, bool):
                        val = "是" if val else "否"
                    label = ("⚠️ " if any(i.field_key == fv.key for i in result.issues) else "") + fv.name
                    new_val = st.text_input(label, value=str(val) if val is not None else "", key=f"edit_{result.file_id}_{fv.key}")
                    edited_vals[fv.key] = new_val or None

                if st.button("💾 应用并保存修正", type="primary", key=f"save_{result.file_id}", use_container_width=True):
                    try:
                        updated = patch_result(
                            st.session_state.task_id,
                            result.file_id,
                            edited_vals,
                        )
                        st.session_state.results = [
                            updated if r.file_id == result.file_id else r
                            for r in st.session_state.results
                        ]
                        st.session_state.edit_mode = False
                        st.toast("✅ 数据已修正，系统已重新演算状态。", icon="💾")
                        time.sleep(0.5)
                        st.rerun()
                    except Exception as e:
                        st.error(f"保存数据库失败：{e}")

    if result.status.value in ("failed", "cancelled"):
        with st.container(border=True):
            st.markdown("💡 此文件因异常被中断，您可以尝试调整左侧边栏的模型参数后重试。")
            if st.button("🔄 重新投入引擎队列", type="primary", key=f"retry_{result.file_id}"):
                try:
                    retry_file(st.session_state.task_id, result.file_id)
                    from task_engine import get_results
                    new_data = get_results(st.session_state.task_id, page_size=500)
                    st.session_state.results = new_data["results"]
                    st.toast("✅ 已重新加入处理队列。", icon="🔄")
                    time.sleep(0.5)
                    st.rerun()
                except Exception as e:
                    st.error(f"重试唤醒失败：{e}")


def _export_bar():
    from schema import ExportScope

    with st.container(border=True):
        st.markdown("<div class='card-title'>📤 将当前批次数据打包导出</div>", unsafe_allow_html=True)
        ec1, ec2, ec3, ec4, ec5 = st.columns(5)

        scope_label = ec1.selectbox("数据选择范围", ["成功+待检查", "全量(含失败)", "仅选择当前文件"], label_visibility="collapsed")
        scope_map   = {"成功+待检查": ExportScope.SUCCESS_ONLY, "全量(含失败)": ExportScope.ALL, "仅选择当前文件": ExportScope.SELECTED}
        scope       = scope_map[scope_label]
        include_log = ec2.checkbox("携带推理追溯日志", value=True)

        if not st.session_state.task_id:
            return

        from task_engine import get_task as _get_task
        from template_service import get_template
        from export_service import build_excel, build_docx
        from schema import ExportRequest

        task     = _get_task(st.session_state.task_id)
        template = get_template(task.template_id)
        req      = ExportRequest(scope=scope, include_log=include_log)

        from app_fastapi import _select_results
        results_to_export = _select_results(task, req)

        excel_key = f"excel_bytes_{task.task_id}"
        docx_key = f"docx_bytes_{task.task_id}"

        with ec3:
            if st.button("📊 汇总为 Excel", use_container_width=True, type="primary", key="gen_xl"):
                with st.spinner("结构化装配 Excel 中…"):
                    st.session_state[excel_key] = build_excel(task, template, results_to_export, include_log=include_log)
            
            if excel_key in st.session_state:
                st.download_button(
                    "📥 获取 .xlsx 文件",
                    data=st.session_state[excel_key],
                    file_name=f"Data_Export_{task.task_id[:8]}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_excel",
                    use_container_width=True,
                )

        with ec4:
            if st.button("📝 拆分为 Word", use_container_width=True, key="gen_docx"):
                with st.spinner("生成排版 DOCX 报告中…"):
                    st.session_state[docx_key] = build_docx(task, template, results_to_export, include_log=include_log)
            
            if docx_key in st.session_state:
                st.download_button(
                    "📥 获取 .docx 压缩包",
                    data=st.session_state[docx_key],
                    file_name=f"Report_{task.task_id[:8]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_docx",
                    use_container_width=True,
                )
        
        with ec5:
            st.caption("✨ 想生成定制化的汇总报告或精美 PPT？请前往『AI 控制台』进行多文件融合交互！")


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — AI CONSOLE (Upgraded with Template Mimicking & Native DOCX)
# ══════════════════════════════════════════════════════════════════════════════

_QUICK_COMMANDS = [
    ("📝 仿写汇总报告",  "请严格参考【参考排版样板】的格式、标题层级与公文语气，将【待处理目标文档】中的议案进行整理汇总。并在末尾补充分析目标文档中缺失的关键要素（如：关联交易缺少公允性描述、缺少资金监管说明等）。"),
    ("📄 智能深度摘要",  "请忽略样板格式，直接对目标文档进行结构化的深度摘要，分为背景、核心发现、和结论。"),
    ("📊 数据指标提取",  "请将目标文档中所有财务数据或数字指标整理为 Markdown 表格格式。"),
    ("✅ 识别执行清单",  "请梳理目标文档中分配给不同团队的所有行动项(To-Do)和建议，并分配优先级。"),
]

def page_console():
    st.markdown("<div class='sec-hdr'>💬 AI 深度分析与仿写控制台 <span style='font-size:0.8rem;font-weight:600;color:#9CA3AF;margin-left:auto;background:#F3F4F6;padding:4px 10px;border-radius:20px;'>Step 4 of 4</span></div>", unsafe_allow_html=True)

    model = st.session_state.get("selected_model", "gpt-4o-mini")

    col_left, col_right = st.columns([1, 1.8], gap="large")

    with col_left:
        with st.container(border=True):
            st.markdown("<div class='card-title'>📂 语料装载枢纽</div>", unsafe_allow_html=True)
            st.markdown("<p style='font-size:0.85rem;color:#6B7280;margin-bottom:16px;'>在此构建您的工作上下文。您可以提供一份标准样板让 AI 学习其公文排版风格，再批量投喂目标源文件进行汇总或审查。</p>", unsafe_allow_html=True)

            # --- Uploader 1: Template Document ---
            def sync_tpl_callback():
                f = st.session_state.get("console_tpl_uploader")
                if f:
                    st.session_state["console_tpl_list"] = [{"name": f.name, "raw": f.getvalue(), "size": f.size}]
                else:
                    st.session_state["console_tpl_list"] = []

            st.markdown("<div style='font-weight:600; color:#4B2E83; margin-bottom:8px;'>🎯 第一步：装载参考样板 (可选)</div>", unsafe_allow_html=True)
            st.file_uploader(
                "上传1份模板文档，AI将克隆其行文风格与大纲排版",
                type=["pdf", "docx", "txt"],
                key="console_tpl_uploader",
                on_change=sync_tpl_callback,
            )
            if st.session_state.get("console_tpl_list"):
                st.caption(f"✅ 已锁定样板文件：{st.session_state['console_tpl_list'][0]['name']}")

            st.markdown("<hr style='margin:16px 0; border:none; border-top:1px dashed #E5E7EB;'>", unsafe_allow_html=True)

            # --- Uploader 2: Target Documents ---
            def sync_tgt_callback():
                raw_uploaded = st.session_state.get("console_tgt_uploader")
                if raw_uploaded:
                    cached = [{"name": f.name, "raw": f.getvalue(), "size": f.size} for f in raw_uploaded]
                    st.session_state["console_file_list"] = cached
                else:
                    st.session_state["console_file_list"] = []

            st.markdown("<div style='font-weight:600; color:#4B2E83; margin-bottom:8px;'>📂 第二步：装载待处理目标文档</div>", unsafe_allow_html=True)
            st.file_uploader(
                "上传需要被汇总或审查的实体文档 (支持多选)",
                type=["pdf", "docx"],
                key="console_tgt_uploader",
                accept_multiple_files=True,
                on_change=sync_tgt_callback,
            )

            file_list = st.session_state.get("console_file_list", [])

            if file_list:
                n = len(file_list)
                total_kb = sum(f["size"] for f in file_list) // 1024
                st.markdown(f"""
                <div style='background:#F0FDF4; border:1px solid #A7F3D0; border-radius:8px; padding:12px; margin:12px 0;'>
                    <div style='color:#065F46; font-weight:700; font-size:0.85rem;'>待处理池已就绪：{n} 份文件 (共 {total_kb} KB)</div>
                </div>
                """, unsafe_allow_html=True)

                col_e, col_c = st.columns([2, 1])
                extract_clicked = col_e.button("⚡ 解析并注入大模型", use_container_width=True, type="primary")
                col_c.button("🗑️ 清空重载", use_container_width=True, on_click=lambda: st.session_state.update({"console_file_list": [], "console_tpl_list": []}))

                if extract_clicked:
                    prog = st.progress(0, text="⏳ 引擎预热中…")
                    try:
                        import concurrent.futures
                        from extractor import extract
                        
                        # 1. 解析样板文件 (如果存在)
                        tpl_text = ""
                        tpl_list = st.session_state.get("console_tpl_list", [])
                        if tpl_list:
                            prog.progress(10, text="🔍 正在拆解并学习参考样板排版…")
                            fitem = tpl_list[0]
                            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                                fut = pool.submit(extract, fitem["name"], fitem["raw"])
                                try:
                                    res = fut.result(timeout=45)
                                    tpl_text = res.full_text if not isinstance(res, str) else res
                                except Exception:
                                    tpl_text = ""
                        
                        # 2. 解析目标文件
                        all_texts = []
                        sep = "\n\n"
                        for fi, fitem in enumerate(file_list):
                            pct = 30 + int(60 * fi / max(len(file_list), 1))
                            prog.progress(pct, text=f"🔍 深度萃取目标文件 {fi + 1}/{len(file_list)}…")
                            
                            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                                fut = pool.submit(extract, fitem["name"], fitem["raw"])
                                try:
                                    txt = fut.result(timeout=45)
                                    if not isinstance(txt, str):
                                        txt = txt.full_text
                                except concurrent.futures.TimeoutError:
                                    txt = f"[{fitem['name']} 解析超时]"
                            all_texts.append(f"【目标文档源: {fitem['name']}】\n{txt}")

                        full_text = sep.join(all_texts)
                        doc_names = " + ".join(f["name"] for f in file_list)

                        st.session_state["chat_template_text"] = tpl_text
                        st.session_state["chat_template_name"] = tpl_list[0]["name"] if tpl_list else ""
                        st.session_state["chat_doc_text"] = full_text
                        st.session_state["chat_doc_name"] = doc_names
                        
                        msg = f"🔗 底层语料链已建立。注入 {len(file_list)} 份目标文件。"
                        if tpl_text:
                            msg += f" 并已成功克隆 1 份格式样板引擎库。"
                            
                        st.session_state["chat_history"].append({"role": "system_info", "content": msg})
                        prog.progress(100, text=f"🎉 上下文装载完毕！")
                        st.toast("语料库挂载成功，您可以开始仿写或对话了！", icon="✨")
                        time.sleep(0.5)
                        st.rerun()

                    except Exception as e:
                        prog.empty()
                        st.error(f"解析崩溃：{e}")

        if st.session_state.get("chat_doc_name"):
            with st.container(border=True):
                st.markdown("<div class='card-title' style='color:#059669;'>🧠 活跃记忆上下文</div>", unsafe_allow_html=True)
                if st.session_state.get("chat_template_name"):
                    st.markdown(f"<div style='font-size:0.85rem; font-weight:600; color:#D97706; margin-bottom:4px;'>🎯 样板锁定: {st.session_state['chat_template_name']}</div>", unsafe_allow_html=True)
                st.markdown(f"<div style='font-size:0.85rem; font-weight:600;'>📂 目标文档: {st.session_state['chat_doc_name']}</div>", unsafe_allow_html=True)
                st.caption(f"驱动引擎: `{model}`")

                if st.button("🔌 熔断并清理记忆区", use_container_width=True, key="clear_extracted_context"):
                    st.session_state["chat_template_text"] = ""
                    st.session_state["chat_template_name"] = ""
                    st.session_state["chat_doc_text"] = ""
                    st.session_state["chat_doc_name"] = ""
                    st.session_state["chat_history"] = []
                    st.rerun()

    with col_right:
        doc_text = st.session_state.get("chat_doc_text", "")
        tpl_text = st.session_state.get("chat_template_text", "")

        if not doc_text:
            st.markdown("""
            <div style="background:#FFFFFF;border:1px dashed #D1D5DB;border-radius:14px;padding:100px 32px;text-align:center;">
              <div style="font-size:3.5rem;margin-bottom:16px;opacity:0.8;">💬</div>
              <div style="font-weight:700;color:#374151;font-size:1.15rem;margin-bottom:8px;">AI 推理大脑正休眠</div>
              <div style="color:#9CA3AF;font-size:0.9rem;line-height:1.6;max-width:380px;margin:0 auto;">
                请在左侧区域分别上传【样板文件】与【待分析文档】。我们将通过大模型进行像素级的公文排版克隆与内容汇总。
              </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            with st.container(border=True):
                chat_container = st.container(height=450)
                with chat_container:
                    if not st.session_state["chat_history"]:
                        st.markdown("<div style='text-align:center; color:#9CA3AF; font-size:0.9rem; padding: 20px;'>您可以从下方挑选快捷指令（如：仿写汇总报告），或直接在输入框提问。</div>", unsafe_allow_html=True)
                    for msg in st.session_state["chat_history"]:
                        if msg["role"] == "user":
                            with st.chat_message("user", avatar="👤"):
                                st.markdown(msg["content"])
                        elif msg["role"] == "assistant":
                            with st.chat_message("assistant", avatar="🤖"):
                                st.markdown(msg["content"])
                        elif msg["role"] == "system_info":
                            st.info(msg["content"])

            # ── Free-form chat input ──────────────────────────────────────
            user_input = st.chat_input("向 AI 下达审查编写指令或询问文档细节...")

            if user_input:
                # ── 核心逻辑：组装双轨 Prompt，执行样板级联克隆 ──
                sys_prompt = "你是极其严谨的商业文档分析顾问与公文写作专家。\n"
                if tpl_text:
                    sys_prompt += f"\n【参考排版样板】\n请严格剖析并完全克隆以下样板文本的大纲层级、行文格式与汇报语气：\n---\n{tpl_text[:6000]}\n---\n\n"
                
                sys_prompt += f"【待处理目标文档】\n请根据以下目标文档的内容进行分析或重组作答：\n---\n{doc_text[:16000]}\n---\n\n"
                
                sys_prompt += (
                    "【严格执行原则】\n"
                    "1. 绝不允许编造或幻觉出文档中不存在的数据和议案。\n"
                    "2. 最终输出请使用清晰的 Markdown 排版。\n"
                )
                if tpl_text:
                    sys_prompt += "3. 用户要求格式化输出或汇总时，必须像模版一样组织你的标题和要点结构。\n"

                last_msgs = st.session_state["chat_history"]
                last_user = next((m for m in reversed(last_msgs) if m["role"] == "user"), None)
                if not last_user or last_user["content"] != user_input:
                    st.session_state["chat_history"].append({"role": "user", "content": user_input})
                    with st.spinner("AI 深度仿写与推理中…"):
                        answer = _call_chat_llm(sys_prompt, user_input, model)
                    st.session_state["chat_history"].append({"role": "assistant", "content": answer})
                st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)
            
            # ── 快捷指令 Pills ──
            with st.container(border=True):
                st.markdown("<div class='card-title' style='margin-bottom:8px; border-bottom:none; padding-bottom:0;'>✨ 深度汇编指令库</div>", unsafe_allow_html=True)
                rows = [_QUICK_COMMANDS[i:i+4] for i in range(0, len(_QUICK_COMMANDS), 4)]
                for row in rows:
                    cols = st.columns(len(row))
                    for col, (label, prompt) in zip(cols, row):
                        if col.button(label, use_container_width=True, key=f"q_{label}"):
                            st.session_state["chat_history"].append({"role": "user", "content": prompt})
                            with st.spinner(f"正在全速执行: {label}…"):
                                sys_p = "你是专业公文写作专家。\n"
                                if tpl_text:
                                    sys_p += f"\n【参考排版样板】\n请严格克隆以下结构和语气：\n---\n{tpl_text[:6000]}\n---\n"
                                sys_p += f"\n【待处理目标文档】\n---\n{doc_text[:16000]}\n---\n"
                                answer = _call_chat_llm(sys_p, prompt, model)
                            st.session_state["chat_history"].append({"role": "assistant", "content": answer})
                            st.rerun()

            # ── 成果物导出 (新增 DOCX 引擎) ──
            last_ai_msg = next((m["content"] for m in reversed(st.session_state["chat_history"]) if m["role"] == "assistant"), None)
            if last_ai_msg:
                with st.expander("📤 将最新对话成果物封装打包", expanded=True):
                    st.markdown("<p style='font-size:0.85rem; color:#6B7280;'>一键将大模型输出的汇总报告或表格导出为本地原生办公文档，直接投入工作流使用。</p>", unsafe_allow_html=True)
                    ec1, ec2, ec3 = st.columns(3)
                    
                    with ec1:
                        if st.button("📝 编译生成 Word 文稿", use_container_width=True, type="secondary"):
                            with st.spinner("正在将 Markdown 引擎接驳至 DOCX 原生排版器..."):
                                st.session_state["ai_docx_gen_bytes"] = _ai_response_to_docx_local(last_ai_msg, st.session_state.get("chat_doc_name", "document"))
                        
                        if "ai_docx_gen_bytes" in st.session_state:
                            st.download_button(
                                "⬇ 下载汇总汇报.docx",
                                data=st.session_state["ai_docx_gen_bytes"],
                                file_name=f"Report_{time.strftime('%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True,
                            )
                            
                    with ec2:
                        if st.button("📊 提取矩阵至 Excel", use_container_width=True, type="secondary"):
                            with st.spinner("正在提取数据为 Excel 矩阵…"):
                                st.session_state["ai_excel_bytes"] = _ai_response_to_excel(last_ai_msg, st.session_state.get("chat_doc_name", "document"))
                        
                        if "ai_excel_bytes" in st.session_state:
                            st.download_button(
                                "⬇ 下载数据表.xlsx",
                                data=st.session_state["ai_excel_bytes"],
                                file_name=f"Data_{time.strftime('%H%M')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                type="primary",
                                use_container_width=True,
                            )
                            
                    with ec3:
                        if st.button("📑 构建结构化 PPT", use_container_width=True, type="secondary"):
                            with st.spinner("🧠 编排 Agent 正在重构幻灯片逻辑流..."):
                                st.session_state["ai_ppt_bytes"] = _ai_response_to_pptx(last_ai_msg, doc_text, st.session_state.get("chat_doc_name", "document"), model)
                        
                        if "ai_ppt_bytes" in st.session_state:
                            st.download_button(
                                "⬇ 下载演示文稿.pptx",
                                data=st.session_state["ai_ppt_bytes"],
                                file_name=f"Presentation_{time.strftime('%H%M')}.pptx",
                                mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                                type="primary",
                                use_container_width=True,
                            )


# ══════════════════════════════════════════════════════════════════════════════
# HELPER LLM CALLS
# ══════════════════════════════════════════════════════════════════════════════

def _call_chat_llm(system_prompt: str, user_msg: str, model: str) -> str:
    import os
    try:
        if model.startswith("claude"):
            from langchain_anthropic import ChatAnthropic
            from langchain_core.messages import HumanMessage, SystemMessage
            llm = ChatAnthropic(model=model, temperature=0.15, max_tokens=3000)
        else:
            from langchain_openai import ChatOpenAI
            from langchain_core.messages import HumanMessage, SystemMessage
            base_url = None
            if model.startswith("qwen"):
                base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
                api_key  = os.environ.get("DASHSCOPE_API_KEY", "")
                if not api_key:
                    return "❌ 未找到 DASHSCOPE_API_KEY。请在侧边栏「🔑 API 密钥」配置。"
                llm = ChatOpenAI(model=model, temperature=0.15, max_tokens=3000, base_url=base_url, api_key=api_key)
            elif model.startswith("deepseek"):
                base_url = "https://api.deepseek.com/v1"
                api_key  = os.environ.get("DEEPSEEK_API_KEY", os.environ.get("OPENAI_API_KEY", ""))
                llm = ChatOpenAI(model=model, temperature=0.15, max_tokens=3000, base_url=base_url, api_key=api_key)
            elif model.startswith("gemini"):
                base_url = "https://generativelanguage.googleapis.com/v1beta/openai/"
                api_key  = os.environ.get("GOOGLE_API_KEY", os.environ.get("OPENAI_API_KEY", ""))
                llm = ChatOpenAI(model=model, temperature=0.15, max_tokens=3000, base_url=base_url, api_key=api_key)
            else:
                llm = ChatOpenAI(model=model, temperature=0.15, max_tokens=3000)

        resp = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=user_msg)])
        return resp.content.strip()
    except Exception as e:
        return f"❌ 对话引擎异常：{e}"


def _parse_md_tables(text):
    tables = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and not line.startswith("|---"):
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            if i + 1 < len(lines) and lines[i+1].strip().startswith("|---"):
                data_rows = []
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    row = [c.strip() for c in lines[j].strip("|").split("|")]
                    data_rows.append(row)
                    j += 1
                tables.append((header_cells, data_rows))
                i = j
                continue
        i += 1
    return tables

def _ai_response_to_docx_local(ai_text: str, doc_name: str) -> bytes:
    """Native conversion from AI Markdown directly to Docx"""
    import io
    import re
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required. Please install it with: pip install python-docx")

    doc = Document()
    
    # Apply global styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # Main Title
    title = doc.add_heading('智能文档审阅与汇总分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title.runs[0].font
    title_font.name = 'Microsoft YaHei'
    title_font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"来源数据列阵: {doc_name}\n编译生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run_meta.font.size = Pt(9)
    run_meta.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()

    # Iterative basic markdown parsing
    for line in ai_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.runs[0].font.name = 'Microsoft YaHei'
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.runs[0].font.name = 'Microsoft YaHei'
            h.runs[0].font.color.rgb = RGBColor(0x4A, 0x2C, 0x99)
        elif line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.runs[0].font.name = 'Microsoft YaHei'
            h.runs[0].font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)
        elif line.startswith('- ') or line.startswith('* '):
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
            p = doc.add_paragraph(clean_text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            clean_text = re.sub(r'^\d+\.\s', '', line)
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            p = doc.add_paragraph(clean_text, style='List Number')
        else:
            # Handle inline bolding basic regex mapping
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            doc.add_paragraph(clean_line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ai_response_to_excel(ai_text, doc_name):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    PURPLE = "4B2E83"
    LIGHT_PUR = "EDE7F6"
    hdr_fill = PatternFill(start_color=PURPLE, end_color=PURPLE, fill_type="solid")
    alt_fill = PatternFill(start_color=LIGHT_PUR, end_color=LIGHT_PUR, fill_type="solid")
    wf = Font(bold=True, color="FFFFFF", size=11)
    thin = Border(left=Side(style="thin", color="CCCCCC"), right=Side(style="thin", color="CCCCCC"), top=Side(style="thin", color="CCCCCC"), bottom=Side(style="thin", color="CCCCCC"))

    def _hdr(ws, row_idx):
        for cell in ws[row_idx]:
            cell.fill = hdr_fill; cell.font = wf
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin

    def _row(ws, row_idx, alt=False):
        for cell in ws[row_idx]:
            if alt: cell.fill = alt_fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin

    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    tables = _parse_md_tables(ai_text)

    if tables:
        for ti, (headers, rows) in enumerate(tables):
            ws = wb.create_sheet(f"数据表格{ti + 1}")
            ws.append(headers)
            _hdr(ws, 1)
            for ri, row in enumerate(rows):
                padded = (row + [""] * len(headers))[:len(headers)]
                ws.append(padded)
                _row(ws, ri + 2, alt=(ri % 2 == 0))
            for col in ws.columns:
                max_w = max((len(str(cell.value or "")) for cell in col), default=10)
                ws.column_dimensions[col[0].column_letter].width = min(max_w + 4, 50)
    else:
        ws_kv = wb.create_sheet("核心摘录")
        ws_kv.append(["序号", "关键句萃取"])
        _hdr(ws_kv, 1)
        items = []
        for line in ai_text.splitlines():
            line = line.strip()
            mm = re.match(r"^[-*\u2022]\s+(.+)$", line) or re.match(r"^\d+[.)]\s+(.+)$", line)
            if mm: items.append(mm.group(1))
        if items:
            for ri, item in enumerate(items):
                parts = re.split(r"[\uff1a:]\s*", item, maxsplit=1)
                ws_kv.append([parts[0].strip(), parts[1].strip()] if len(parts) == 2 else [str(ri+1), item])
                _row(ws_kv, ri + 2, alt=(ri % 2 == 0))
        else:
            for ri, s in enumerate(filter(None, re.split(r"[。！？.!?]\s*", ai_text))):
                ws_kv.append([str(ri+1), s.strip()])
                _row(ws_kv, ri + 2, alt=(ri % 2 == 0))
        ws_kv.column_dimensions["A"].width = 8
        ws_kv.column_dimensions["B"].width = 90

    ws_raw = wb.create_sheet("分析原档")
    ws_raw["A1"] = "AI 推理流原档"
    ws_raw["A1"].font = Font(bold=True, size=13, color=PURPLE)
    for ln, line in enumerate(ai_text.splitlines(), start=3):
        ws_raw.cell(row=ln, column=1, value=line)
    ws_raw.column_dimensions["A"].width = 120

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _ai_response_to_pptx(ai_text: str, doc_text: str, doc_name: str, model: str) -> bytes:
    from export_service import build_pptx
    import json
    import re
    
    system_prompt = """你是一个高级 PPT 数据构造引擎。提取文档内容，输出符合以下 JSON Schema 的纯净 JSON 字符串，不要带 markdown 代码块。
{
    "summary": "核心执行摘要", "topics": ["主题1"], "entities": ["实体1"], "dates": ["日期"], "actions": ["行动"],
    "page_count": 1, "char_count": 2000,
    "sections": [
        {"title": "财务目标", "content": [["指标", "数值"], ["净利润", "100万"]], "type": "table"}
    ]
}"""
    raw_response = _call_chat_llm(system_prompt, f"AI思考:{ai_text}\n原文:{doc_text[:10000]}", model)
    try:
        # 完全规避由于硬编码三个反引号导致的任何 Markdown 解析器截断或崩溃 Bug
        MD_TICK = chr(96) * 3
        cleaned = re.sub(r"^" + MD_TICK + r"(?:json)?\s*", "", raw_response.strip())
        cleaned = re.sub(r"\s*" + MD_TICK + r"$", "", cleaned).strip()
        info = json.loads(cleaned)
    except Exception as e:
        info = {
            "summary": "解析失败回退", 
            "sections": [{"title": "提取异常", "content": f"无法解析 JSON。错误: {str(e)}", "type": "text"}]
        }
    return build_pptx(info, source_filename=doc_name)


# ══════════════════════════════════════════════════════════════════════════════
# ROUTER
# ══════════════════════════════════════════════════════════════════════════════

_sidebar()
page = st.session_state.page
if page == "1_template": page_template()
elif page == "2_extract": page_extract()
elif page == "3_review": page_review()
elif page == "4_console": page_console()
else: page_template()