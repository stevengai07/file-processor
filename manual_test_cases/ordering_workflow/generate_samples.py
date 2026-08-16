"""Generate small sample files for manual Streamlit ordering workflow checks."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook


BASE_DIR = Path(__file__).resolve().parent
DOCS_DIR = BASE_DIR / "docs"


def create_ordering_excel() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "发言顺序"
    ws.append(["顺序号", "时间", "发言标题", "地区"])
    rows = [
        [1, "09:00-09:10", "人工智能产业发展报告", "北京"],
        [2, "09:10-09:20", "数字政府建设汇报", "上海"],
        [3, "09:20-09:30", "绿色能源项目材料", "广东"],
        [4, "09:30-09:40", "智慧交通建设方案", "浙江"],
        [5, "09:40-09:50", "未提交材料测试", "江苏"],
    ]
    for row in rows:
        ws.append(row)
    for column, width in {"A": 10, "B": 18, "C": 28, "D": 12}.items():
        ws.column_dimensions[column].width = width
    wb.save(BASE_DIR / "ordering_sample.xlsx")


def create_docx(filename: str, lines: list[str]) -> None:
    doc = Document()
    doc.add_heading(lines[0].replace("标题：", ""), level=1)
    for line in lines:
        doc.add_paragraph(line)
    doc.save(DOCS_DIR / filename)


def create_text_pdf(filename: str, lines: list[str]) -> None:
    # Minimal Type 0 CID font PDF so Chinese text can be stored as extractable UTF-16BE strings.
    text_ops = []
    y = 760
    for line in lines:
        hex_text = line.encode("utf-16-be").hex().upper()
        text_ops.append(f"BT /F1 12 Tf 72 {y} Td <{hex_text}> Tj ET")
        y -= 22
    stream = "\n".join(text_ops).encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 7 0 R >>",
        b"<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light /Encoding /UniGB-UCS2-H /DescendantFonts [5 0 R] >>",
        b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light /CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 2 >> /FontDescriptor 6 0 R >>",
        b"<< /Type /FontDescriptor /FontName /STSong-Light /Flags 6 /FontBBox [0 -200 1000 900] /ItalicAngle 0 /Ascent 800 /Descent -200 /CapHeight 700 /StemV 80 >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f\n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n\n".encode("ascii"))
    content.extend(
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    (DOCS_DIR / filename).write_bytes(bytes(content))


def create_corrupt_pdf() -> None:
    (DOCS_DIR / "unreadable_corrupt.pdf").write_bytes(
        b"%PDF-1.4\nThis file is intentionally incomplete for failure isolation testing."
    )


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    create_ordering_excel()
    create_text_pdf(
        "人工智能产业发展报告.pdf",
        [
            "标题：人工智能产业发展报告",
            "发言单位：北京市发展研究中心",
            "本文介绍人工智能产业发展现状、政策支持、重点企业和未来规划。",
        ],
    )
    create_docx(
        "数字政府建设汇报.docx",
        [
            "标题：数字政府建设汇报",
            "发言单位：上海市大数据中心",
            "本文汇报数字政府平台建设、数据共享、政务服务优化等内容。",
        ],
    )
    create_text_pdf(
        "绿色能源项目材料.pdf",
        [
            "标题：绿色能源项目材料",
            "发言单位：广东省能源集团",
            "本文介绍绿色能源项目建设进展、投资规模和后续计划。",
        ],
    )
    create_docx(
        "重复候选_人工智能产业发展报告.docx",
        [
            "标题：人工智能产业发展报告",
            "发言单位：备用材料单位",
            "这是一个重复候选文件，用于验证人工复核或候选冲突展示。",
        ],
    )
    create_text_pdf(
        "unrelated_无关材料.pdf",
        [
            "标题：无关材料",
            "发言单位：测试单位",
            "该文件不应匹配任何顺序项，用于验证未使用文件统计。",
        ],
    )
    create_corrupt_pdf()


if __name__ == "__main__":
    main()
