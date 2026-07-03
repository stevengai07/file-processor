#!/usr/bin/env python3
"""
extractor.py — Document text extraction + OCR pipeline.

Responsibilities:
  1. Accept raw file bytes + filename.
  2. Dispatch to the correct reader (PDF or DOCX).
  3. For PDFs: attempt native text extraction first (pdfplumber);
     if the page is image-only (low char density), fall back to OCR.
  4. For DOCX: extract paragraphs, tables, headers, footers, and text boxes.
  5. Apply image enhancement before Tesseract (delegates to image_enhancer.py).
  6. Return an ExtractedDocument with per-page text and metadata.

Public API:
  extract(filename, raw, settings) -> ExtractedDocument

Dependencies:
  pip install pdfplumber python-docx pytesseract pillow opencv-python-headless
"""

from __future__ import annotations

import io
import re
import time
import unicodedata
from dataclasses import dataclass, field
from typing import List, Optional, Tuple, Dict, Any


# ══════════════════════════════════════════════════════════════════════════════
# Data classes
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class PageText:
    """Holds extracted text and metadata for a single page."""
    page_num: int
    text: str
    char_count: int = 0
    used_ocr: bool = False
    ocr_confidence: float = 0.0
    ocr_psm: int = 0            # Tesseract PSM mode actually used
    word_count: int = 0

    def __post_init__(self):
        stripped = self.text.strip()
        self.char_count = len(stripped)
        self.word_count = len(stripped.split()) if stripped else 0

    def is_empty(self) -> bool:
        return self.char_count < 5

    def __repr__(self) -> str:
        ocr_flag = " [OCR]" if self.used_ocr else ""
        return (f"<PageText p={self.page_num}{ocr_flag} "
                f"chars={self.char_count} words={self.word_count}>")


@dataclass
class ExtractedDocument:
    """
    Result of extracting a full document.

    Attributes:
        pages        : list of PageText, one per page (or logical section for DOCX)
        metadata     : dict of document metadata (author, title, created, etc.)
        filename     : original filename
        elapsed_sec  : extraction wall-clock time in seconds
    """
    pages: List[PageText] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    filename: str = ""
    elapsed_sec: float = 0.0

    # ── Computed properties ──────────────────────────────────────────────────

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def ocr_page_count(self) -> int:
        return sum(1 for p in self.pages if p.used_ocr)

    @property
    def total_chars(self) -> int:
        return sum(p.char_count for p in self.pages)

    @property
    def total_words(self) -> int:
        return sum(p.word_count for p in self.pages)

    @property
    def avg_ocr_confidence(self) -> float:
        ocr_pages = [p for p in self.pages if p.used_ocr]
        if not ocr_pages:
            return 0.0
        return sum(p.ocr_confidence for p in ocr_pages) / len(ocr_pages)

    @property
    def full_text(self) -> str:
        """All pages joined with page markers."""
        parts = []
        for p in self.pages:
            if not p.is_empty():
                parts.append(f"--- Page {p.page_num} ---")
                parts.append(p.text.strip())
        return "\n\n".join(parts)

    @property
    def plain_text(self) -> str:
        """Full text without page markers."""
        return "\n\n".join(p.text.strip() for p in self.pages if not p.is_empty())

    # ── Query helpers ────────────────────────────────────────────────────────

    def text_around(self, keyword: str, window: int = 300) -> str:
        """
        Return up to `window` chars of context around first occurrence of keyword.
        Used by the AI agent to feed targeted excerpts to the LLM.
        """
        haystack = self.full_text
        idx = haystack.lower().find(keyword.lower())
        if idx == -1:
            return ""
        start = max(0, idx - window)
        end = min(len(haystack), idx + len(keyword) + window)
        return haystack[start:end]

    def search(self, keyword: str) -> List[Tuple[int, str]]:
        """
        Return all (page_num, excerpt) tuples where keyword appears.
        Case-insensitive. Returns at most 3 excerpts per page.
        """
        results = []
        kw_lower = keyword.lower()
        for page in self.pages:
            haystack = page.text
            hay_lower = haystack.lower()
            start = 0
            hits = 0
            while hits < 3:
                idx = hay_lower.find(kw_lower, start)
                if idx == -1:
                    break
                snippet_start = max(0, idx - 120)
                snippet_end   = min(len(haystack), idx + len(keyword) + 120)
                results.append((page.page_num, haystack[snippet_start:snippet_end]))
                start = idx + 1
                hits += 1
        return results

    def get_page(self, page_num: int) -> Optional[PageText]:
        """Return a specific page by 1-based page number."""
        for p in self.pages:
            if p.page_num == page_num:
                return p
        return None

    def summary_stats(self) -> dict:
        """Return a dict of summary statistics suitable for logging or display."""
        return {
            "filename":          self.filename,
            "page_count":        self.page_count,
            "total_chars":       self.total_chars,
            "total_words":       self.total_words,
            "ocr_page_count":    self.ocr_page_count,
            "avg_ocr_confidence": round(self.avg_ocr_confidence, 1),
            "elapsed_sec":       round(self.elapsed_sec, 2),
            "metadata":          self.metadata,
        }

    def __repr__(self) -> str:
        return (f"<ExtractedDocument '{self.filename}' "
                f"pages={self.page_count} chars={self.total_chars} "
                f"ocr={self.ocr_page_count}>")


# ══════════════════════════════════════════════════════════════════════════════
# Text cleaning
# ══════════════════════════════════════════════════════════════════════════════

# Ligature and typography normalisation table
_LIGATURES: Dict[str, str] = {
    "\ufb00": "ff",  "\ufb01": "fi",  "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "ft",
    "\u2018": "'",   "\u2019": "'",
    "\u201c": '"',   "\u201d": '"',
    "\u2013": "-",   "\u2014": "--",
    "\u00a0": " ",   "\u200b": "",   "\ufeff": "",
}

# Common OCR digit/letter confusion pairs (applied only to OCR pages)
_OCR_FIXES = [
    (re.compile(r"(?<=[A-Z])0(?=[A-Z])"),   "O"),   # letter-0-letter → O
    (re.compile(r"(?<=\d)O(?=\d)"),         "0"),   # digit-O-digit → 0
    (re.compile(r"(?<=\d)l(?=\d)"),         "1"),   # digit-l-digit → 1
    (re.compile(r"\bl\b"),                  "1"),   # lone "l" → 1 in numeric context
]

# CJK Unicode block ranges (for CJK-specific post-processing)
_CJK_RANGES = [
    (0x4E00, 0x9FFF),   # CJK Unified Ideographs
    (0x3400, 0x4DBF),   # CJK Extension A
    (0x20000, 0x2A6DF), # CJK Extension B
    (0x2A700, 0x2B73F), # CJK Extension C
    (0x2B740, 0x2B81F), # CJK Extension D
    (0xF900, 0xFAFF),   # CJK Compatibility Ideographs
    (0xFE30, 0xFE4F),   # CJK Compatibility Forms
    (0x3000, 0x303F),   # CJK Symbols and Punctuation
    (0xFF00, 0xFFEF),   # Halfwidth/Fullwidth Forms
]

def _is_cjk_char(ch: str) -> bool:
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in _CJK_RANGES)

def _has_cjk(text: str) -> bool:
    return any(_is_cjk_char(c) for c in text)

def _dehyphenate(text: str) -> str:
    """
    Join words broken across lines with hyphens.
    e.g. "evalu-\nation" → "evaluation"
    Skips CJK-dominant text where hyphens may be meaningful.
    """
    if _has_cjk(text):
        return text
    return re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)

def _rejoin_cjk_spaces(text: str) -> str:
    """
    Remove spurious spaces inserted between CJK characters by some OCR engines.
    e.g. "中 文 字 符" → "中文字符"
    """
    def _strip_between(m: re.Match) -> str:
        a, b = m.group(1), m.group(2)
        if _is_cjk_char(a) and _is_cjk_char(b):
            return a + b
        return m.group(0)
    return re.sub(r"(.)\s(.)", _strip_between, text)

def _clean_text(text: str, is_ocr: bool = False) -> str:
    """
    Master text cleaner. Applied to every page after extraction.

    Steps:
      1. Ligature expansion
      2. Null bytes + control chars stripped
      3. Normalize unicode to NFC
      4. NBSP → space
      5. Dehyphenation (Latin text)
      6. CJK space rejoining (OCR output)
      7. OCR digit/letter confusion fixes (OCR pages only)
      8. Collapse 3+ consecutive newlines → 2
      9. Strip trailing whitespace per line
    """
    if not text:
        return ""

    # 1. Ligatures
    for src, dst in _LIGATURES.items():
        text = text.replace(src, dst)

    # 2. Control characters (keep \n \t \r)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # 3. NFC normalisation
    text = unicodedata.normalize("NFC", text)

    # 4. NBSP already covered by ligature table, but ensure
    text = text.replace("\u00a0", " ").replace("\u200b", "")

    # 5. Dehyphenation
    text = _dehyphenate(text)

    # 6. CJK space fix
    if _has_cjk(text):
        text = _rejoin_cjk_spaces(text)

    # 7. OCR-specific fixes
    if is_ocr:
        for pattern, replacement in _OCR_FIXES:
            text = pattern.sub(replacement, text)

    # 8. Collapse excess newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 9. Strip trailing whitespace per line
    lines = [ln.rstrip() for ln in text.splitlines()]
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# OCR helper
# ══════════════════════════════════════════════════════════════════════════════

# Tesseract PSM modes to try, in priority order, based on document type preset
_PSM_BY_PRESET = {
    "scanner":  [6, 3],    # uniform block of text, then auto
    "photo":    [11, 6],   # sparse text, then uniform block
    "mixed":    [3, 6],    # auto, then uniform block
    "default":  [6],
}


def _run_tesseract(
    img,                        # PIL Image
    lang: str,
    psm: int,
    min_confidence: int = 30,
) -> Tuple[str, float, int]:
    """
    Run Tesseract on a PIL image with the given PSM mode.

    Returns:
        (text, avg_confidence, actual_psm)
    Only words with confidence > min_confidence are kept.
    """
    try:
        import pytesseract
    except ImportError:
        raise ImportError("pytesseract required: pip install pytesseract")

    config = f"--oem 1 --psm {psm} -l {lang}"
    data = pytesseract.image_to_data(
        img, config=config,
        output_type=pytesseract.Output.DICT
    )

    words, confs = [], []
    for word, conf in zip(data["text"], data["conf"]):
        c = int(conf) if str(conf).lstrip("-").isdigit() else -1
        if c >= min_confidence and word.strip():
            words.append(word)
            confs.append(c)

    text = " ".join(words)
    avg_conf = sum(confs) / len(confs) if confs else 0.0
    return text, avg_conf, psm


def _ocr_image(img, settings: dict) -> Tuple[str, float, int]:
    """
    Try each PSM mode from the preset; return the result with most characters.

    Args:
        img      : PIL Image (already enhanced by image_enhancer)
        settings : extraction settings dict

    Returns:
        (best_text, best_confidence, best_psm)
    """
    lang   = settings.get("tesseract_lang", "chi_sim+eng")
    preset = settings.get("ocr_preset", "default")
    psm_list = _PSM_BY_PRESET.get(preset, _PSM_BY_PRESET["default"])

    best_text, best_conf, best_psm = "", 0.0, psm_list[0]
    for psm in psm_list:
        text, conf, actual_psm = _run_tesseract(img, lang, psm)
        if len(text.strip()) > len(best_text.strip()):
            best_text, best_conf, best_psm = text, conf, actual_psm

    return best_text, best_conf, best_psm


# ══════════════════════════════════════════════════════════════════════════════
# PDF extraction
# ══════════════════════════════════════════════════════════════════════════════

def _extract_pdf(raw: bytes, settings: dict) -> Tuple[List[PageText], dict]:
    """
    Two-pass PDF extraction strategy:

    Pass 1 — pdfplumber native text
        Fast and 100% accurate for digitally-created PDFs.
        Extracts text, preserving word order via bounding box sort.

    Pass 2 — Tesseract OCR (triggered per-page when char_count < OCR_THRESHOLD)
        Used for scanned pages or image-only PDFs.
        Image is enhanced by image_enhancer before being passed to Tesseract.

    Merge rule:
        If OCR produces ≥ 120% of native char count → use OCR text.
        Otherwise → use native text (avoids degrading quality on mixed pages).

    Args:
        raw      : PDF file bytes
        settings : extraction settings dict

    Returns:
        (pages, metadata)
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber required: pip install pdfplumber")

    try:
        from image_enhancer import enhance as _enhance
    except ImportError:
        def _enhance(img, **kw):
            return img

    OCR_THRESHOLD = settings.get("ocr_threshold", 50)
    OCR_MERGE_RATIO = settings.get("ocr_merge_ratio", 1.2)
    dpi = settings.get("ocr_dpi", 300)
    max_pages = settings.get("max_pages", None)

    pages: List[PageText] = []
    metadata: dict = {}

    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        # Extract metadata
        if pdf.metadata:
            metadata = {
                k: str(v) for k, v in pdf.metadata.items()
                if isinstance(v, (str, int, float, bool))
            }

        total = len(pdf.pages)
        limit = min(total, max_pages) if max_pages else total

        for i, page in enumerate(pdf.pages[:limit], start=1):
            # ── Pass 1: native text ──────────────────────────────────────────
            raw_native = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            native = _clean_text(raw_native)
            native_chars = len(native.strip())

            if native_chars >= OCR_THRESHOLD:
                pages.append(PageText(page_num=i, text=native))
                continue

            # ── Pass 2: OCR ──────────────────────────────────────────────────
            try:
                pil_img = page.to_image(resolution=dpi).original
                pil_img = _enhance(
                    pil_img,
                    deskew=settings.get("ocr_deskew", True),
                    denoise=settings.get("ocr_denoise", True),
                    contrast=settings.get("ocr_contrast", True),
                    binarize=settings.get("ocr_binarize", False),
                )
                ocr_raw, confidence, used_psm = _ocr_image(pil_img, settings)
                ocr_text = _clean_text(ocr_raw, is_ocr=True)
                ocr_chars = len(ocr_text.strip())

                # Merge: use OCR only if it produced substantially more text
                if ocr_chars > native_chars * OCR_MERGE_RATIO:
                    pages.append(PageText(
                        page_num=i, text=ocr_text,
                        used_ocr=True, ocr_confidence=confidence,
                        ocr_psm=used_psm,
                    ))
                else:
                    pages.append(PageText(page_num=i, text=native))

            except Exception as exc:
                # OCR failed — fall back to whatever native text we have
                fallback = native if native.strip() else f"[OCR failed: {exc}]"
                pages.append(PageText(page_num=i, text=fallback))

    return pages, metadata


# ══════════════════════════════════════════════════════════════════════════════
# DOCX extraction
# ══════════════════════════════════════════════════════════════════════════════

def _iter_paragraphs_with_breaks(doc):
    """
    Yield (paragraph, is_page_break) pairs, detecting explicit page breaks
    inside any paragraph run via w:br type="page".
    """
    try:
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    for para in doc.paragraphs:
        has_break = any(
            br.get(qn("w:type")) == "page"
            for br in para._element.iter(qn("w:br"))
        )
        yield para, has_break


def _extract_table(table) -> str:
    """
    Extract a DOCX table to pipe-delimited text.
    Merged cells are de-duplicated within each row.
    """
    seen_row_texts = set()
    rows_text = []

    for row in table.rows:
        seen_in_row: set = set()
        cells = []
        for cell in row.cells:
            ct = cell.text.strip()
            if ct and ct not in seen_in_row:
                seen_in_row.add(ct)
                cells.append(ct)
        if cells:
            row_text = " | ".join(cells)
            if row_text not in seen_row_texts:
                seen_row_texts.add(row_text)
                rows_text.append(row_text)

    return "\n".join(rows_text)


def _extract_text_boxes(doc_element) -> List[str]:
    """
    Extract text from Word text boxes (wps:txbx XML nodes).
    These are often missed by the standard paragraph iterator.
    """
    ns_wps = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    texts = []
    for txbx in doc_element.iter(f"{{{ns_wps}}}txbx"):
        parts = [t.text for t in txbx.iter() if t.text and t.text.strip()]
        if parts:
            texts.append(" ".join(parts))
    return texts


def _extract_header_footer(section) -> List[str]:
    """Extract text from a section's header and footer (skips linked sections)."""
    texts = []
    for hf in (section.header, section.footer):
        try:
            if hf and not hf.is_linked_to_previous:
                t = " ".join(p.text.strip() for p in hf.paragraphs if p.text.strip())
                if t:
                    texts.append(t)
        except Exception:
            pass
    return texts


# Heading style → markdown prefix map
_HEADING_PREFIX: Dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Title":     "# ",
    "Subtitle":  "## ",
}


def _extract_docx(raw: bytes, settings: dict) -> Tuple[List[PageText], dict]:
    """
    DOCX reading-order extraction pipeline:

    1. Body paragraphs    — styled headings preserved as # / ## / ### markers
    2. Page breaks        — on w:br type="page" → flush current page buffer
    3. Tables             — cell-by-cell in row order, pipe-delimited
    4. Text boxes/shapes  — wps:txbx XML nodes
    5. Headers & footers  — per-section, skips is_linked_to_previous
    6. Metadata           — core_properties (author, title, created, modified)

    Pages are split on explicit page breaks. If no page breaks exist, the
    whole document is treated as a single logical page.

    Args:
        raw      : DOCX file bytes
        settings : extraction settings dict

    Returns:
        (pages, metadata)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = Document(io.BytesIO(raw))

    current_page: List[str] = []
    page_chunks:  List[str] = []

    def _flush():
        joined = _clean_text("\n".join(current_page))
        if joined.strip():
            page_chunks.append(joined)
        current_page.clear()

    # ── 1+2. Body paragraphs with page-break detection ────────────────────────
    for para, is_break in _iter_paragraphs_with_breaks(doc):
        if is_break:
            _flush()

        style_name = para.style.name if para.style else ""
        prefix = _HEADING_PREFIX.get(style_name, "")
        text = para.text.strip()

        if text:
            current_page.append(prefix + text)

    # ── 3. Tables ─────────────────────────────────────────────────────────────
    for table in doc.tables:
        table_text = _extract_table(table)
        if table_text:
            current_page.append(table_text)

    # ── 4. Text boxes ─────────────────────────────────────────────────────────
    for tb_text in _extract_text_boxes(doc.element):
        current_page.append(tb_text)

    _flush()  # flush final page

    # ── 5. Headers & footers (append to first page as context) ────────────────
    hf_texts = []
    for section in doc.sections:
        hf_texts.extend(_extract_header_footer(section))
    if hf_texts and page_chunks:
        page_chunks[0] = "\n".join(hf_texts) + "\n\n" + page_chunks[0]

    # ── 6. Metadata ───────────────────────────────────────────────────────────
    cp = doc.core_properties
    metadata = {
        "author":   cp.author   or "",
        "title":    cp.title    or "",
        "subject":  cp.subject  or "",
        "keywords": cp.keywords or "",
        "created":  str(cp.created)  if cp.created  else "",
        "modified": str(cp.modified) if cp.modified else "",
        "revision": str(cp.revision) if cp.revision else "",
        "language": cp.language or "",
    }

    # ── Build PageText list ───────────────────────────────────────────────────
    if not page_chunks:
        page_chunks = [""]  # ensure at least one page

    pages = [
        PageText(page_num=i + 1, text=chunk)
        for i, chunk in enumerate(page_chunks)
        if chunk.strip()
    ]
    if not pages:
        pages = [PageText(page_num=1, text="")]

    return pages, metadata


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

_SUPPORTED_EXTENSIONS = {"pdf", "docx", "doc"}


def extract(
    filename: str,
    raw: bytes,
    settings: dict | None = None,
) -> ExtractedDocument:
    """
    Extract text from a PDF or DOCX file.

    Args:
        filename : original filename — used to determine file type
        raw      : raw file bytes
        settings : optional config overrides (all have defaults):

            # PDF + OCR options
            ocr_threshold   int   50       min native chars before OCR is tried
            ocr_merge_ratio float 1.2      OCR must be 1.2× longer than native to win
            ocr_dpi         int   300      DPI for rendering PDF pages to images
            ocr_deskew      bool  True     deskew images before OCR
            ocr_denoise     bool  True     denoise images before OCR
            ocr_contrast    bool  True     auto-contrast images before OCR
            ocr_binarize    bool  False    adaptive binarize (good for scanned docs)
            ocr_preset      str   "default" PSM preset: scanner/photo/mixed/default
            tesseract_lang  str   "chi_sim+eng"  Tesseract language pack(s)
            max_pages       int   None     limit pages extracted (None = all)

    Returns:
        ExtractedDocument with .pages, .full_text, .metadata,
        .total_chars, .ocr_page_count, .avg_ocr_confidence, etc.

    Raises:
        ValueError  : unsupported file extension
        ImportError : required library not installed
    """
    if settings is None:
        settings = {}

    t_start = time.monotonic()
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext not in _SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: .{ext!r} — "
            f"supported: {sorted(_SUPPORTED_EXTENSIONS)}"
        )

    doc = ExtractedDocument(filename=filename)

    if ext == "pdf":
        doc.pages, doc.metadata = _extract_pdf(raw, settings)
    else:  # docx / doc
        doc.pages, doc.metadata = _extract_docx(raw, settings)

    doc.elapsed_sec = time.monotonic() - t_start
    return doc


# ══════════════════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════════════════

def _cli():
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <file.pdf|file.docx> [--max-pages N] [--lang LANG]")
        print()
        print("Options:")
        print("  --max-pages N    Only extract first N pages")
        print("  --lang LANG      Tesseract language (default: chi_sim+eng)")
        print("  --preset PRESET  OCR preset: scanner|photo|mixed|default")
        print("  --binarize       Enable adaptive binarization before OCR")
        sys.exit(1)

    path = sys.argv[1]
    settings: dict = {}

    args = sys.argv[2:]
    i = 0
    while i < len(args):
        if args[i] == "--max-pages" and i + 1 < len(args):
            settings["max_pages"] = int(args[i + 1]); i += 2
        elif args[i] == "--lang" and i + 1 < len(args):
            settings["tesseract_lang"] = args[i + 1]; i += 2
        elif args[i] == "--preset" and i + 1 < len(args):
            settings["ocr_preset"] = args[i + 1]; i += 2
        elif args[i] == "--binarize":
            settings["ocr_binarize"] = True; i += 1
        else:
            i += 1

    with open(path, "rb") as f:
        raw = f.read()

    result = extract(path, raw, settings)
    stats = result.summary_stats()

    print(f"File       : {stats['filename']}")
    print(f"Pages      : {stats['page_count']}")
    print(f"Total chars: {stats['total_chars']:,}")
    print(f"Total words: {stats['total_words']:,}")
    print(f"OCR pages  : {stats['ocr_page_count']}")
    if stats["ocr_page_count"] > 0:
        print(f"OCR conf.  : {stats['avg_ocr_confidence']:.1f}%")
    print(f"Elapsed    : {stats['elapsed_sec']:.2f}s")
    if stats["metadata"]:
        print(f"Metadata   : {stats['metadata']}")
    print()
    print("── First 800 chars ──")
    print(result.full_text[:800])


if __name__ == "__main__":
    _cli()